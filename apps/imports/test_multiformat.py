import tempfile
from decimal import Decimal
from io import BytesIO

import openpyxl
from django.contrib.auth.models import Group, Permission
from django.core.files.uploadedfile import SimpleUploadedFile
from django.test import TestCase, override_settings
from django.urls import reverse
from docx import Document
from openpyxl.drawing.image import Image as ExcelImage
from PIL import Image as PILImage
from reportlab.pdfgen import canvas

from apps.accounts.models import User
from apps.catalog.models import Categoria, Marca, Producto, Proveedor, UnidadMedida
from apps.pricing.models import HistorialCosto

from .forms import NuevaImportacionForm
from .models import ImportacionFila, ImportacionImagen, ImportacionListaPrecios
from .parsing import analizar_archivo, tipo_archivo_por_nombre
from .services import confirmar_importacion, procesar_importacion, reclasificar_fila


_MEDIA_ROOT = tempfile.mkdtemp(prefix="arqclima-11h-media-")


def _usuario_admin(username="diego_11h"):
    grupo, _ = Group.objects.get_or_create(name="Administrador")
    for codename, app in (
        ("view_historialcosto", "pricing"),
        ("add_historialcosto", "pricing"),
        ("add_marca", "catalog"),
        ("add_producto", "catalog"),
        ("change_producto", "catalog"),
        ("add_proveedor", "catalog"),
    ):
        grupo.permissions.add(
            Permission.objects.get(codename=codename, content_type__app_label=app)
        )
    usuario = User.objects.create_user(username=username, password="clave12345")
    usuario.groups.add(grupo)
    return usuario


def _imagen_png_bytes(color=(80, 120, 160), size=(240, 160)):
    buffer = BytesIO()
    PILImage.new("RGB", size, color).save(buffer, format="PNG")
    return buffer.getvalue()


def _xlsx_con_imagen():
    libro = openpyxl.Workbook()
    hoja = libro.active
    hoja.title = "Precios"
    hoja.append(["Marca", "Código", "Descripción", "Precio Neto", "Unidad", "Categoría"])
    hoja.append(["Vulcano", "VX-100", "Válvula VX", "1.234,50", "unidad", "Válvulas"])

    imagen_stream = BytesIO(_imagen_png_bytes())
    imagen = ExcelImage(imagen_stream)
    imagen.width = 120
    imagen.height = 80
    hoja.add_image(imagen, "G2")

    buffer = BytesIO()
    libro.save(buffer)
    buffer.seek(0)
    return SimpleUploadedFile(
        "lista-con-imagen.xlsx",
        buffer.read(),
        content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )




def _xlsx_lista_comercial_por_bloques():
    """
    Replica el patrón de listas reales de proveedor donde la cabecera
    nombra categoría/embalaje/precios pero NO rotula Código ni Nombre.
    """
    libro = openpyxl.Workbook()
    hoja = libro.active
    hoja.title = "LISTA GENERAL"

    hoja["C5"] = "CODO A 90º HH"
    hoja["E5"] = "EMBALAJE"
    hoja["F5"] = "PRECIO $"
    hoja["G5"] = "PRECIO BONIF"
    hoja["H5"] = "PEDIDO"
    hoja["I5"] = "SUB TOTAL"

    hoja["C7"] = 101203105200000
    hoja["D7"] = "CODO A 90º Ø 20"
    hoja["E7"] = 200
    hoja["F7"] = 100
    hoja["G7"] = "$ 50.00"

    hoja["C8"] = 101203105250000
    hoja["D8"] = "CODO A 90º Ø 25"
    hoja["E8"] = 120
    hoja["F8"] = 200
    hoja["G8"] = "$ 100.00"

    # Producto anunciado pero todavía sin precio efectivo.
    hoja["C9"] = 101203105400000
    hoja["D9"] = "CODO A 90º Ø 40 - PROXIMAMENTE"
    hoja["F9"] = 0
    hoja["G9"] = "$ -"

    hoja["C13"] = "CUPLA HH"
    hoja["E13"] = "EMBALAJE"
    hoja["F13"] = "PRECIO $"
    hoja["G13"] = "PRECIO BONIF"
    hoja["H13"] = "PEDIDO"
    hoja["I13"] = "SUB TOTAL"

    hoja["C15"] = "101203108200001"
    hoja["D15"] = "CUPLA UNION Ø 20"
    hoja["E15"] = 300
    hoja["F15"] = 80
    hoja["G15"] = "$ 40.00"

    buffer = BytesIO()
    libro.save(buffer)
    buffer.seek(0)
    return SimpleUploadedFile(
        "lista-comercial-bloques.xlsx",
        buffer.read(),
        content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )


def _docx_con_tabla_e_imagen():
    documento = Document()
    tabla = documento.add_table(rows=2, cols=5)
    encabezados = ["Marca", "Código", "Nombre", "Costo", "Código proveedor"]
    valores = ["Vulcano", "DOC-1", "Producto Word", "850,25", "WORD-77"]
    for indice, texto in enumerate(encabezados):
        tabla.rows[0].cells[indice].text = texto
    for indice, texto in enumerate(valores):
        tabla.rows[1].cells[indice].text = texto

    run = documento.add_paragraph().add_run()
    run.add_picture(BytesIO(_imagen_png_bytes((160, 90, 70))), width=None)

    buffer = BytesIO()
    documento.save(buffer)
    buffer.seek(0)
    return SimpleUploadedFile(
        "lista.docx",
        buffer.read(),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


def _pdf_tabular():
    buffer = BytesIO()
    pdf = canvas.Canvas(buffer)
    y = 800
    posiciones = (40, 150, 270, 470)
    for x, texto in zip(posiciones, ("Marca", "Código", "Nombre", "Costo")):
        pdf.drawString(x, y, texto)
    y -= 24
    for x, texto in zip(posiciones, ("Vulcano", "PDF-1", "Producto PDF", "990,50")):
        pdf.drawString(x, y, texto)
    pdf.save()
    buffer.seek(0)
    return SimpleUploadedFile("lista.pdf", buffer.read(), content_type="application/pdf")


def _pdf_escaneado():
    imagen = PILImage.new("RGB", (800, 600), (245, 245, 245))
    buffer = BytesIO()
    imagen.save(buffer, format="PDF")
    buffer.seek(0)
    return SimpleUploadedFile("escaneado.pdf", buffer.read(), content_type="application/pdf")


@override_settings(MEDIA_ROOT=_MEDIA_ROOT)
class AnalisisMultiformatoTests(TestCase):
    def setUp(self):
        self.usuario = _usuario_admin()
        self.proveedor = Proveedor.objects.create(nombre_comercial="Proveedor 11H")
        self.marca = Marca.objects.create(nombre="Vulcano")
        self.categoria = Categoria.objects.create(nombre="Válvulas")

    def _crear_importacion(self, archivo):
        tipo = tipo_archivo_por_nombre(archivo.name)
        importacion = ImportacionListaPrecios.objects.create(
            proveedor=self.proveedor,
            archivo=archivo,
            tipo_archivo=tipo,
            cargado_por=self.usuario,
        )
        procesar_importacion(importacion)
        importacion.refresh_from_db()
        return importacion

    def test_csv_detecta_columnas_y_producto_nuevo(self):
        archivo = SimpleUploadedFile(
            "lista.csv",
            "Marca;Código;Nombre;Costo;Unidad\nVulcano;CSV-1;Producto CSV;1500,75;unidad\n".encode(
                "utf-8"
            ),
            content_type="text/csv",
        )

        importacion = self._crear_importacion(archivo)

        fila = importacion.filas.get()
        self.assertEqual(importacion.tipo_archivo, ImportacionListaPrecios.TipoArchivo.CSV)
        self.assertEqual(fila.codigo, "CSV-1")
        self.assertEqual(fila.costo, Decimal("1500.75"))
        self.assertEqual(fila.categoria, ImportacionFila.Categoria.NUEVO_PRODUCTO)
        self.assertEqual(fila.confianza, ImportacionFila.Confianza.ALTA)

    def test_excel_extrae_fila_e_imagen_embebida(self):
        importacion = self._crear_importacion(_xlsx_con_imagen())

        fila = importacion.filas.get()
        imagen = importacion.imagenes.get()

        self.assertEqual(fila.origen, "Hoja Precios")
        self.assertEqual(fila.costo, Decimal("1234.50"))
        self.assertEqual(fila.unidad_texto, "unidad")
        self.assertEqual(fila.categoria_texto, "Válvulas")
        self.assertEqual(imagen.numero_fila_origen, 2)
        self.assertGreater(imagen.ancho, 0)
        self.assertGreater(imagen.alto, 0)



    def test_excel_por_bloques_detecta_codigo_nombre_y_precio_bonificado(self):
        importacion = self._crear_importacion(_xlsx_lista_comercial_por_bloques())

        filas = list(importacion.filas.order_by("numero_fila"))
        self.assertEqual(len(filas), 4)

        primera = filas[0]
        self.assertEqual(primera.codigo, "101203105200000")
        self.assertEqual(primera.nombre_texto, "CODO A 90º Ø 20")
        self.assertEqual(primera.costo, Decimal("50.00"))
        self.assertEqual(primera.categoria_texto, "CODO A 90º HH")
        self.assertEqual(primera.categoria, ImportacionFila.Categoria.PARA_REVISAR)
        self.assertFalse(primera.incluir)

        sin_precio = next(fila for fila in filas if fila.numero_fila == 9)
        self.assertEqual(sin_precio.costo, Decimal("0.00"))
        self.assertEqual(sin_precio.categoria, ImportacionFila.Categoria.ERROR)

        self.assertTrue(
            any(
                "bloque(s) comerciales" in advertencia
                for advertencia in importacion.advertencias_analisis
            )
        )

    def test_excel_por_bloques_no_confunde_embalaje_con_codigo(self):
        importacion = self._crear_importacion(_xlsx_lista_comercial_por_bloques())

        codigos = set(importacion.filas.values_list("codigo", flat=True))
        self.assertIn("101203105200000", codigos)
        self.assertIn("101203108200001", codigos)
        self.assertNotIn("200", codigos)
        self.assertNotIn("300", codigos)


    def test_word_extrae_tabla_e_imagen(self):
        importacion = self._crear_importacion(_docx_con_tabla_e_imagen())

        fila = importacion.filas.get()

        self.assertEqual(fila.origen, "Tabla 1 de Word")
        self.assertEqual(fila.codigo, "DOC-1")
        self.assertEqual(fila.codigo_proveedor_texto, "WORD-77")
        self.assertEqual(fila.costo, Decimal("850.25"))
        self.assertGreaterEqual(importacion.imagenes.count(), 1)

    def test_pdf_con_texto_detecta_tabla_con_confianza_media(self):
        importacion = self._crear_importacion(_pdf_tabular())

        fila = importacion.filas.get()

        self.assertEqual(fila.codigo, "PDF-1")
        self.assertEqual(fila.confianza, ImportacionFila.Confianza.MEDIA)
        self.assertFalse(fila.incluir)
        self.assertEqual(
            importacion.estado_analisis,
            ImportacionListaPrecios.EstadoAnalisis.REQUIERE_REVISION,
        )

    def test_pdf_escaneado_no_hace_ocr_y_conserva_imagen(self):
        importacion = self._crear_importacion(_pdf_escaneado())

        self.assertFalse(importacion.filas.exists())
        self.assertGreaterEqual(importacion.imagenes.count(), 1)
        self.assertEqual(
            importacion.estado_analisis,
            ImportacionListaPrecios.EstadoAnalisis.REQUIERE_REVISION,
        )
        self.assertTrue(
            any("OCR" in advertencia for advertencia in importacion.advertencias_analisis)
        )

    def test_categoria_y_unidad_existentes_se_aplican_al_producto_nuevo(self):
        archivo = SimpleUploadedFile(
            "producto.csv",
            "Marca;Código;Nombre;Costo;Unidad;Categoría\n"
            "Vulcano;CAT-1;Producto categorizado;100;m;Válvulas\n".encode("utf-8"),
            content_type="text/csv",
        )
        importacion = self._crear_importacion(archivo)

        confirmar_importacion(importacion, self.usuario)

        producto = Producto.objects.get(marca=self.marca, codigo="CAT-1")
        self.assertEqual(producto.categoria, self.categoria)
        self.assertEqual(producto.unidad_medida, UnidadMedida.METRO)
        self.assertTrue(
            HistorialCosto.objects.filter(
                producto_proveedor__producto=producto,
                costo=Decimal("100.00"),
            ).exists()
        )


@override_settings(MEDIA_ROOT=_MEDIA_ROOT)
class RevisionSeguraImportacionTests(TestCase):
    def setUp(self):
        self.usuario = _usuario_admin("diego_revision_11h")
        self.proveedor = Proveedor.objects.create(nombre_comercial="Proveedor revisión 11H")

    def _importacion_csv(self, texto):
        archivo = SimpleUploadedFile(
            "revision.csv",
            texto.encode("utf-8"),
            content_type="text/csv",
        )
        importacion = ImportacionListaPrecios.objects.create(
            proveedor=self.proveedor,
            archivo=archivo,
            tipo_archivo=ImportacionListaPrecios.TipoArchivo.CSV,
            cargado_por=self.usuario,
        )
        procesar_importacion(importacion)
        return importacion

    def test_fila_sin_marca_queda_para_revisar_y_no_se_puede_forzar(self):
        importacion = self._importacion_csv(
            "Código;Nombre;Costo\nSIN-MARCA;Producto sin marca;100\n"
        )
        fila = importacion.filas.get()

        self.assertEqual(fila.categoria, ImportacionFila.Categoria.PARA_REVISAR)
        self.assertFalse(fila.incluir)

        fila.incluir = True
        fila.save(update_fields=["incluir"])
        resultado = confirmar_importacion(importacion, self.usuario)

        self.assertEqual(resultado["creados"], 0)
        self.assertEqual(resultado["omitidos"], 1)
        self.assertFalse(Producto.objects.filter(codigo="SIN-MARCA").exists())

    def test_editar_fila_reclasifica_y_marca_como_revisada(self):
        importacion = self._importacion_csv(
            "Código;Nombre;Costo\nEDIT-1;Producto editado;200\n"
        )
        fila = importacion.filas.get()

        reclasificar_fila(
            fila,
            self.usuario,
            {
                "marca": "Marca Corregida",
                "codigo": "EDIT-1",
                "nombre": "Producto editado",
                "descripcion": "",
                "costo_crudo": Decimal("200"),
                "codigo_proveedor": "",
                "unidad": "unidad",
                "categoria": "",
            },
        )
        fila.refresh_from_db()

        self.assertEqual(fila.confianza, ImportacionFila.Confianza.REVISADA)
        self.assertEqual(fila.categoria, ImportacionFila.Categoria.NUEVO_PRODUCTO)
        self.assertTrue(fila.incluir)

    def test_duplicados_con_costos_distintos_bloquean_ambas_filas(self):
        importacion = self._importacion_csv(
            "Marca;Código;Nombre;Costo\n"
            "Vulcano;DUP-1;Duplicado;100\n"
            "Vulcano;DUP-1;Duplicado;120\n"
        )

        self.assertEqual(
            importacion.filas.filter(categoria=ImportacionFila.Categoria.PARA_REVISAR).count(),
            2,
        )
        self.assertFalse(importacion.filas.filter(incluir=True).exists())

    def test_duplicado_identico_se_aplica_una_sola_vez(self):
        importacion = self._importacion_csv(
            "Marca;Código;Nombre;Costo\n"
            "Vulcano;DUP-2;Duplicado idéntico;100\n"
            "Vulcano;DUP-2;Duplicado idéntico;100\n"
        )

        filas = list(importacion.filas.order_by("pk"))
        self.assertTrue(filas[0].incluir)
        self.assertFalse(filas[1].incluir)
        self.assertEqual(filas[1].categoria, ImportacionFila.Categoria.SIN_CAMBIOS)


@override_settings(MEDIA_ROOT=_MEDIA_ROOT)
class VistasMultiformatoTests(TestCase):
    def setUp(self):
        self.usuario = _usuario_admin("diego_vistas_11h")
        self.proveedor = Proveedor.objects.create(nombre_comercial="Proveedor vistas 11H")
        self.client.login(username=self.usuario.username, password="clave12345")



    def test_nueva_importacion_permite_crear_proveedor_sin_salir(self):
        response = self.client.get(reverse("imports:nueva"))

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "Crear proveedor")
        self.assertContains(
            response,
            reverse("catalog:proveedor_nuevo_rapido"),
        )
        self.assertContains(response, 'id="proveedor-rapido-form"')

    def test_alta_rapida_crea_proveedor_activo_y_devuelve_json(self):
        response = self.client.post(
            reverse("catalog:proveedor_nuevo_rapido"),
            {
                "nombre_comercial": "Proveedor Nuevo 11H",
                "razon_social": "Proveedor Nuevo SA",
                "cuit": "30-11111111-1",
                "contacto_nombre": "Compras",
                "telefono": "2610000000",
                "email": "compras@proveedor.test",
            },
            HTTP_X_REQUESTED_WITH="XMLHttpRequest",
        )

        self.assertEqual(response.status_code, 201)
        data = response.json()
        self.assertTrue(data["ok"])
        proveedor = Proveedor.objects.get(pk=data["proveedor"]["id"])
        self.assertTrue(proveedor.activo)
        self.assertEqual(proveedor.nombre_comercial, "Proveedor Nuevo 11H")
        self.assertEqual(data["proveedor"]["email"], "compras@proveedor.test")

    def test_alta_rapida_rechaza_nombre_o_cuit_duplicados(self):
        existente = Proveedor.objects.create(
            nombre_comercial="Duplicado 11H",
            cuit="30-22222222-2",
        )

        por_nombre = self.client.post(
            reverse("catalog:proveedor_nuevo_rapido"),
            {
                "nombre_comercial": existente.nombre_comercial.lower(),
                "cuit": "30-33333333-3",
            },
        )
        self.assertEqual(por_nombre.status_code, 400)
        self.assertIn("nombre_comercial", por_nombre.json()["errores"])

        por_cuit = self.client.post(
            reverse("catalog:proveedor_nuevo_rapido"),
            {
                "nombre_comercial": "Otro proveedor",
                "cuit": existente.cuit,
            },
        )
        self.assertEqual(por_cuit.status_code, 400)
        self.assertIn("cuit", por_cuit.json()["errores"])

    def test_usuario_sin_add_proveedor_no_ve_ni_usa_alta_rapida(self):
        grupo, _ = Group.objects.get_or_create(name="Importador sin alta proveedor")
        permiso_importar = Permission.objects.get(
            codename="add_importacionlistaprecios",
            content_type__app_label="imports",
        )
        grupo.permissions.add(permiso_importar)
        usuario = User.objects.create_user(
            username="importador_sin_proveedor_11h",
            password="clave12345",
        )
        usuario.groups.add(grupo)

        self.client.logout()
        self.client.login(
            username=usuario.username,
            password="clave12345",
        )

        pagina = self.client.get(reverse("imports:nueva"))
        self.assertEqual(pagina.status_code, 200)
        self.assertNotContains(pagina, "Crear proveedor")

        alta = self.client.post(
            reverse("catalog:proveedor_nuevo_rapido"),
            {"nombre_comercial": "No permitido"},
        )
        self.assertEqual(alta.status_code, 403)
        self.assertFalse(
            Proveedor.objects.filter(nombre_comercial="No permitido").exists()
        )

    def test_formulario_acepta_formatos_nuevos_y_rechaza_doc_antiguo(self):
        form = NuevaImportacionForm(
            data={"proveedor": self.proveedor.pk},
            files={
                "archivo": SimpleUploadedFile(
                    "lista.doc",
                    b"contenido",
                    content_type="application/msword",
                )
            },
        )
        self.assertFalse(form.is_valid())
        self.assertIn("Formato no soportado", str(form.errors["archivo"]))

    def test_subir_csv_redirige_a_preview_multiformato(self):
        archivo = SimpleUploadedFile(
            "lista.csv",
            b"Marca;Codigo;Nombre;Costo\nVulcano;V-1;Valvula;100\n",
            content_type="text/csv",
        )

        response = self.client.post(
            reverse("imports:nueva"),
            {"proveedor": self.proveedor.pk, "archivo": archivo},
        )

        importacion = ImportacionListaPrecios.objects.get()
        self.assertRedirects(response, reverse("imports:detalle", args=[importacion.pk]))
        self.assertEqual(importacion.tipo_archivo, ImportacionListaPrecios.TipoArchivo.CSV)

    def test_preview_muestra_imagenes_y_archivo_por_rutas_protegidas(self):
        archivo = _xlsx_con_imagen()
        importacion = ImportacionListaPrecios.objects.create(
            proveedor=self.proveedor,
            archivo=archivo,
            tipo_archivo=ImportacionListaPrecios.TipoArchivo.XLSX,
            cargado_por=self.usuario,
        )
        procesar_importacion(importacion)
        imagen = importacion.imagenes.get()

        detalle = self.client.get(reverse("imports:detalle", args=[importacion.pk]))
        self.assertEqual(detalle.status_code, 200)
        self.assertContains(detalle, "Imágenes detectadas")
        self.assertContains(detalle, reverse("imports:imagen", args=[importacion.pk, imagen.pk]))

        respuesta_imagen = self.client.get(
            reverse("imports:imagen", args=[importacion.pk, imagen.pk])
        )
        self.assertEqual(respuesta_imagen.status_code, 200)

        respuesta_archivo = self.client.get(reverse("imports:archivo", args=[importacion.pk]))
        self.assertEqual(respuesta_archivo.status_code, 200)

    def test_edicion_desde_preview_corrige_fila(self):
        archivo = SimpleUploadedFile(
            "sin-marca.csv",
            b"Codigo;Nombre;Costo\nEDIT-VISTA;Producto;100\n",
            content_type="text/csv",
        )
        importacion = ImportacionListaPrecios.objects.create(
            proveedor=self.proveedor,
            archivo=archivo,
            tipo_archivo=ImportacionListaPrecios.TipoArchivo.CSV,
            cargado_por=self.usuario,
        )
        procesar_importacion(importacion)
        fila = importacion.filas.get()

        response = self.client.post(
            reverse("imports:fila_editar", args=[importacion.pk, fila.pk]),
            {
                "marca": "Marca Vista",
                "codigo": "EDIT-VISTA",
                "nombre": "Producto",
                "descripcion": "",
                "costo": "100.00",
                "codigo_proveedor": "",
                "unidad": "unidad",
                "categoria": "",
            },
        )

        self.assertRedirects(response, reverse("imports:detalle", args=[importacion.pk]))
        fila.refresh_from_db()
        self.assertEqual(fila.categoria, ImportacionFila.Categoria.NUEVO_PRODUCTO)
        self.assertEqual(fila.confianza, ImportacionFila.Confianza.REVISADA)



    def test_marca_masiva_reclasifica_lista_sin_columna_marca(self):
        marca = Marca.objects.create(nombre="Tubofusion")
        archivo = _xlsx_lista_comercial_por_bloques()
        importacion = ImportacionListaPrecios.objects.create(
            proveedor=self.proveedor,
            archivo=archivo,
            tipo_archivo=ImportacionListaPrecios.TipoArchivo.XLSX,
            cargado_por=self.usuario,
        )
        procesar_importacion(importacion)

        self.assertEqual(importacion.filas.filter(marca_texto="").count(), 4)

        response = self.client.post(
            reverse("imports:asignar_marca", args=[importacion.pk]),
            {"marca": marca.pk},
        )

        self.assertRedirects(response, reverse("imports:detalle", args=[importacion.pk]))
        self.assertFalse(importacion.filas.filter(marca_texto="").exists())
        self.assertEqual(
            importacion.filas.filter(
                categoria=ImportacionFila.Categoria.NUEVO_PRODUCTO,
                incluir=True,
            ).count(),
            3,
        )
        self.assertEqual(
            importacion.filas.filter(categoria=ImportacionFila.Categoria.ERROR).count(),
            1,
        )

    def test_preview_ofrece_marca_masiva_cuando_falta_columna_marca(self):
        archivo = _xlsx_lista_comercial_por_bloques()
        importacion = ImportacionListaPrecios.objects.create(
            proveedor=self.proveedor,
            archivo=archivo,
            tipo_archivo=ImportacionListaPrecios.TipoArchivo.XLSX,
            cargado_por=self.usuario,
        )
        procesar_importacion(importacion)

        response = self.client.get(reverse("imports:detalle", args=[importacion.pk]))

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, "La planilla no informa Marca")
        self.assertContains(response, "Detectamos códigos, nombres y precios")
        self.assertContains(response, reverse("imports:asignar_marca", args=[importacion.pk]))


    def test_confirmacion_por_post_no_admite_fila_para_revisar_manipulada(self):
        archivo = SimpleUploadedFile(
            "forzada.csv",
            b"Codigo;Nombre;Costo\nFORZADA;Producto;100\n",
            content_type="text/csv",
        )
        importacion = ImportacionListaPrecios.objects.create(
            proveedor=self.proveedor,
            archivo=archivo,
            tipo_archivo=ImportacionListaPrecios.TipoArchivo.CSV,
            cargado_por=self.usuario,
        )
        procesar_importacion(importacion)
        fila = importacion.filas.get()

        response = self.client.post(
            reverse("imports:confirmar", args=[importacion.pk]),
            {"incluir": [str(fila.pk)]},
        )

        self.assertRedirects(response, reverse("imports:detalle", args=[importacion.pk]))
        self.assertFalse(Producto.objects.filter(codigo="FORZADA").exists())
