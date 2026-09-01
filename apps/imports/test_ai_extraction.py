"""
Tests de la Etapa 11 (IA en Imports): mapeo de columnas con Haiku, cache
ProveedorColumnMapping, detección de Excel "tipo catálogo" y extracción
completa con Sonnet (PDF escaneado, imagen, docx sin tabla). Nunca se le
pega a la API real: todo lo que habla con Claude vive en apps.imports.ai,
así que acá se mockea services.mapear_columnas / services.extraer_lista_precios
(los nombres tal como quedan importados en services.py, no en ai.py).
"""

import tempfile
from io import BytesIO
from pathlib import Path
from unittest.mock import MagicMock, patch

import anthropic
import openpyxl
from django.contrib.auth.models import Group, Permission
from django.core.files.uploadedfile import SimpleUploadedFile
from django.test import TestCase, override_settings
from docx import Document
from PIL import Image as PILImage
from reportlab.pdfgen import canvas

from apps.accounts.models import User
from apps.catalog.models import Marca, Proveedor
from apps.pricing.models import Moneda

from .ai import ExtraccionIAError, extraer_lista_precios, mapear_columnas
from .models import ImportacionFila, ImportacionListaPrecios, ProveedorColumnMapping
from .services import procesar_importacion

_MEDIA_ROOT = tempfile.mkdtemp(prefix="arqclima-11-ia-media-")
_FIXTURES_DIR = Path(__file__).parent / "test_fixtures"


def _usuario_admin(username="diego_ia"):
    grupo, _ = Group.objects.get_or_create(name="Administrador")
    for codename, app in (
        ("view_historialcosto", "pricing"),
        ("add_historialcosto", "pricing"),
        ("add_marca", "catalog"),
        ("add_producto", "catalog"),
        ("change_producto", "catalog"),
    ):
        grupo.permissions.add(
            Permission.objects.get(codename=codename, content_type__app_label=app)
        )
    usuario = User.objects.create_user(username=username, password="clave12345")
    usuario.groups.add(grupo)
    return usuario


def _xlsx_bytes(filas):
    libro = openpyxl.Workbook()
    hoja = libro.active
    for fila in filas:
        hoja.append(fila)
    buffer = BytesIO()
    libro.save(buffer)
    buffer.seek(0)
    return buffer.read()


def _xlsx_plano_con_headers_no_reconocidos():
    """
    Encabezado donde código/nombre/costo no matchean ningún alias conocido
    (así que ni el clásico ni el fallback por bloques —que también busca
    una columna de precio por alias/prefijo— encuentran nada), pero "Marca"
    sí matchea de forma exacta: alcanza para que _mejor_candidata_encabezado
    lo elija como candidato y dispare el camino de mapeo con Haiku, sin que
    la columna de costo "Valor Total" dispare por sí sola el fallback por
    bloques (que busca específicamente una columna de precio).
    """
    return SimpleUploadedFile(
        "lista-headers-raros.xlsx",
        _xlsx_bytes(
            [
                ["Marca", "Ident. Producto", "Detalle Comercial", "Valor Total"],
                ["Vulcano", "ABC-123", "Tubo Fusion 20mm", "1.500,00"],
            ]
        ),
        content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )


def _xlsx_catalogo_headers_repetidos():
    """
    Réplica del caso real reportado: dos subtablas en la misma hoja, cada
    una con su propio header Código/Nombre/Precio, separadas por blancos y
    ruido. _contar_filas_header debería detectar 2 headers independientes.
    """
    return SimpleUploadedFile(
        "lista-catalogo.xlsx",
        _xlsx_bytes(
            [
                ["Cliente:", "Distribuidora XYZ"],
                [],
                ["TUBO FUSION PN12"],
                ["Código", "Nombre", "Precio"],
                ["TF-001", "Tubo Fusion PN12 20mm", "1500,00"],
                [],
                ["TUBO FUSION PN20"],
                ["Código", "Nombre", "Precio"],
                ["TF-002", "Tubo Fusion PN20 20mm *NUEVO*", "1800,00"],
            ]
        ),
        content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )


def _xlsx_producto_real_sin_encabezado_ni_bloques():
    """
    Filas con pinta real de producto (código+nombre+precio con formato de
    moneda) pero SIN ningún header reconocible en toda la hoja (ni clásico
    ni por bloques: ninguna celda es literalmente "Código"/"Nombre"/
    "Precio" ni ningún alias conocido) — el caso genuino que sí debe
    escalar a Sonnet por el último recurso de
    _analizar_matriz_con_posible_ia, y que el pre-filtro nuevo (ver
    _fila_parece_producto_confiable) no debe frenar.
    """
    return SimpleUploadedFile(
        "lista-sin-encabezado.xlsx",
        _xlsx_bytes(
            [
                ["Lista informal, sin columnas rotuladas"],
                ["TF-9001", "Tubo Fusion PN12 20mm reforzado", "$ 1.850,00"],
                ["TF-9002", "Codo 90 fusion 20mm", "$ 640,00"],
                ["TF-9003", "Valvula esferica 1/2 pulgada", "1.200,00"],
            ]
        ),
        content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )


def _xlsx_extracto_tipo_mara_sin_precio():
    """
    Réplica reducida (60 filas en vez de las ~9600 reales) del extracto SAP
    "MARA" del Cotizador Base real que motivó el pre-filtro: sin ninguna
    columna de precio, pero con columnas numéricas (ID de material, peso)
    y una descripción con medidas embebidas ("...2.00 m") que antes
    colaban como "costo" para parsear_costo()/_fila_parece_producto(), sin
    que ningún header ni bloque de columnas fuera detectable.

    Va acompañada de una segunda hoja con datos de producto normales
    (header clásico Código/Nombre/Costo) para que la importación tenga
    algo real que resolver — igual que el archivo real tiene, además de
    "MARA", otras hojas ("Cotizador", "GRILLA CLIENTE") que sí se
    resuelven local. Así el test aísla el comportamiento de la hoja tipo
    MARA sin depender del caso borde de "todo el archivo vino vacío"
    (procesar_importacion() lo trata como ColumnasNoDetectadas, algo
    ortogonal a lo que este test verifica).
    """
    libro = openpyxl.Workbook()
    hoja_mara = libro.active
    hoja_mara.title = "MARA"
    for indice in range(60):
        hoja_mara.append(
            [
                10_000_000 + indice,
                f"GA{indice:06d}",
                f"77980811{indice:05d}",
                "C/U",
                0.2,
                0.18,
                f"TUBO COAXIAL O60/100 {(indice % 5) + 1}.00 m",
            ]
        )
    hoja_normal = libro.create_sheet("Cotizador")
    hoja_normal.append(["Código", "Nombre", "Costo"])
    hoja_normal.append(["CAL-100", "Caldera Premium 24kW", "150000,00"])

    buffer = BytesIO()
    libro.save(buffer)
    buffer.seek(0)
    return SimpleUploadedFile(
        "extracto-mara.xlsx",
        buffer.read(),
        content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )


def _pdf_tabular():
    buffer = BytesIO()
    pdf = canvas.Canvas(buffer)
    y = 800
    posiciones = (40, 150, 270, 470)
    for x, texto in zip(posiciones, ("Marca", "Código", "Nombre", "Costo")):
        pdf.drawString(x, y, texto)
    y -= 24
    for x, texto in zip(posiciones, ("Vulcano", "PDF-1", "Producto PDF", "990,50")):
        pdf.drawString(x, y, texto)
    pdf.save()
    buffer.seek(0)
    return SimpleUploadedFile("lista.pdf", buffer.read(), content_type="application/pdf")


def _pdf_escaneado():
    imagen = PILImage.new("RGB", (800, 600), (245, 245, 245))
    buffer = BytesIO()
    imagen.save(buffer, format="PDF")
    buffer.seek(0)
    return SimpleUploadedFile("escaneado.pdf", buffer.read(), content_type="application/pdf")


def _docx_sin_tabla_util():
    documento = Document()
    documento.add_paragraph("Lista de precios variada, sin ninguna tabla ni columnas tabuladas.")
    documento.add_paragraph("Consultar por WhatsApp para más info.")
    buffer = BytesIO()
    documento.save(buffer)
    buffer.seek(0)
    return SimpleUploadedFile(
        "lista-sin-tabla.docx",
        buffer.read(),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


def _imagen_jpg():
    buffer = BytesIO()
    PILImage.new("RGB", (200, 120), (30, 60, 90)).save(buffer, format="JPEG")
    buffer.seek(0)
    return SimpleUploadedFile("foto-lista.jpg", buffer.read(), content_type="image/jpeg")


@override_settings(MEDIA_ROOT=_MEDIA_ROOT, ANTHROPIC_API_KEY="sk-ant-test-clave-secreta")
class MapeoColumnasIATests(TestCase):
    def setUp(self):
        self.usuario = _usuario_admin()
        self.proveedor = Proveedor.objects.create(nombre_comercial="Proveedor IA", activo=True)
        Marca.objects.create(nombre="Vulcano")

    def _importar(self, archivo):
        importacion = ImportacionListaPrecios.objects.create(
            proveedor=self.proveedor,
            archivo=archivo,
            tipo_archivo="xlsx",
            cargado_por=self.usuario,
        )
        procesar_importacion(importacion)
        return importacion

    @patch("apps.imports.services.mapear_columnas")
    def test_cache_miss_llama_a_haiku_y_cachea_el_mapeo(self, mock_mapear):
        mock_mapear.return_value = (
            {"marca": 0, "codigo": 1, "nombre": 2, "costo": 3},
            [],
            {"modelo": "haiku-test", "input_tokens": 10, "output_tokens": 5},
        )

        importacion = self._importar(_xlsx_plano_con_headers_no_reconocidos())

        mock_mapear.assert_called_once()
        self.assertEqual(ProveedorColumnMapping.objects.count(), 1)
        self.assertTrue(importacion.usa_ia)

        fila = importacion.filas.get()
        self.assertEqual(fila.marca_texto, "Vulcano")
        self.assertEqual(fila.codigo, "ABC-123")
        self.assertEqual(fila.nombre_texto, "Tubo Fusion 20mm")
        self.assertEqual(fila.confianza, ImportacionFila.Confianza.ALTA)

    @patch("apps.imports.services.mapear_columnas")
    def test_cache_hit_no_vuelve_a_llamar_a_la_api(self, mock_mapear):
        mock_mapear.return_value = (
            {"marca": 0, "codigo": 1, "nombre": 2, "costo": 3},
            [],
            {"modelo": "haiku-test", "input_tokens": 10, "output_tokens": 5},
        )
        self._importar(_xlsx_plano_con_headers_no_reconocidos())
        self.assertEqual(mock_mapear.call_count, 1)

        # Segunda importación, mismo proveedor, mismos encabezados.
        self._importar(_xlsx_plano_con_headers_no_reconocidos())
        self.assertEqual(mock_mapear.call_count, 1)
        self.assertEqual(ProveedorColumnMapping.objects.count(), 1)

    @patch("apps.imports.services.mapear_columnas")
    def test_respuesta_de_ia_incompleta_no_se_aplica(self, mock_mapear):
        # Falta "costo": _obtener_mapeo_columnas debe rechazarlo.
        mock_mapear.return_value = (
            {"marca": 0, "codigo": 1, "nombre": 2},
            [],
            {"modelo": "haiku-test"},
        )

        importacion = ImportacionListaPrecios.objects.create(
            proveedor=self.proveedor,
            archivo=_xlsx_plano_con_headers_no_reconocidos(),
            tipo_archivo="xlsx",
            cargado_por=self.usuario,
        )
        from .parsing import ColumnasNoDetectadas

        with self.assertRaises(ColumnasNoDetectadas):
            procesar_importacion(importacion)

        self.assertEqual(ProveedorColumnMapping.objects.count(), 0)
        self.assertFalse(importacion.filas.exists())


@override_settings(MEDIA_ROOT=_MEDIA_ROOT, ANTHROPIC_API_KEY="sk-ant-test-clave-secreta")
class ExtraccionCompletaIATests(TestCase):
    def setUp(self):
        self.usuario = _usuario_admin()
        self.proveedor = Proveedor.objects.create(nombre_comercial="Proveedor IA", activo=True)

    def _importar(self, archivo, tipo_archivo):
        importacion = ImportacionListaPrecios.objects.create(
            proveedor=self.proveedor,
            archivo=archivo,
            tipo_archivo=tipo_archivo,
            cargado_por=self.usuario,
        )
        procesar_importacion(importacion)
        return importacion

    @patch("apps.imports.services.extraer_lista_precios")
    def test_excel_catalogo_con_headers_repetidos_escala_a_sonnet(self, mock_extraer):
        mock_extraer.return_value = (
            [
                {
                    "marca": "", "codigo": "TF-001", "nombre": "Tubo Fusion PN12 20mm",
                    "descripcion": "", "costo_texto": "1500,00", "codigo_proveedor": "",
                    "unidad": "", "categoria": "TUBO FUSION PN12", "nota": "",
                },
                {
                    "marca": "", "codigo": "TF-002", "nombre": "Tubo Fusion PN20 20mm",
                    "descripcion": "", "costo_texto": "1800,00", "codigo_proveedor": "",
                    "unidad": "", "categoria": "TUBO FUSION PN20",
                    "nota": "*NUEVO* según el proveedor.",
                },
            ],
            [],
            {"modelo": "sonnet-test", "input_tokens": 500, "output_tokens": 80},
        )

        importacion = self._importar(_xlsx_catalogo_headers_repetidos(), "xlsx")

        mock_extraer.assert_called_once()
        _, kwargs = mock_extraer.call_args
        self.assertEqual(mock_extraer.call_args.args[1], "excel_celdas")

        self.assertEqual(importacion.filas.count(), 2)
        for fila in importacion.filas.all():
            self.assertEqual(fila.confianza, ImportacionFila.Confianza.MEDIA)
            self.assertIn("(IA)", fila.origen)
        fila_nueva = importacion.filas.get(codigo="TF-002")
        self.assertIn("NUEVO", fila_nueva.detalle)
        self.assertTrue(importacion.usa_ia)

    @patch("apps.imports.services.extraer_lista_precios")
    def test_producto_real_sin_encabezado_sigue_escalando_a_sonnet(self, mock_extraer):
        """
        Regresión del pre-filtro agregado para la hoja MARA: datos con
        pinta real de producto (código+nombre+precio con formato de
        moneda), sin header ni bloques detectables, tienen que seguir
        llegando al último recurso de _analizar_matriz_con_posible_ia y
        escalar a Sonnet como antes. El pre-filtro nuevo solo debe frenar
        hojas donde NINGUNA fila tiene esa pinta (ver el test de la hoja
        tipo MARA), no este caso.
        """
        mock_extraer.return_value = (
            [
                {
                    "marca": "", "codigo": "TF-9001", "nombre": "Tubo Fusion PN12 20mm reforzado",
                    "descripcion": "", "costo_texto": "1850,00", "codigo_proveedor": "",
                    "unidad": "", "categoria": "", "nota": "",
                },
            ],
            [],
            {"modelo": "sonnet-test", "input_tokens": 300, "output_tokens": 60},
        )

        importacion = self._importar(_xlsx_producto_real_sin_encabezado_ni_bloques(), "xlsx")

        mock_extraer.assert_called_once()
        self.assertEqual(mock_extraer.call_args.args[1], "excel_celdas")
        self.assertTrue(importacion.usa_ia)

    @patch("apps.imports.services.extraer_lista_precios")
    def test_extracto_tipo_mara_sin_columna_de_precio_no_llama_a_sonnet(self, mock_extraer):
        """
        Caso real que motivó el pre-filtro: una hoja tipo extracto SAP, sin
        ninguna columna de precio, no debe gastar una llamada a Sonnet solo
        para que la IA concluya "esto no tiene precios" — parsear_costo()
        ya alcanza para descartarlo localmente. Antes de este pre-filtro,
        la hoja MARA real (~9600 filas) escalaba igual porque
        parsear_costo() acepta cualquier número de Excel (peso, ID, EAN) y
        una medida embebida en la descripción ("...2.00 m") también
        matcheaba como "costo".
        """
        importacion = self._importar(_xlsx_extracto_tipo_mara_sin_precio(), "xlsx")

        mock_extraer.assert_not_called()
        self.assertFalse(importacion.usa_ia)
        # Solo la fila de la hoja "Cotizador" (resuelta local, camino
        # clásico) — la hoja "MARA" no aportó ninguna.
        self.assertEqual(importacion.filas.count(), 1)
        self.assertEqual(importacion.filas.get().codigo, "CAL-100")
        self.assertTrue(
            any(
                "no se detectó ninguna fila con estructura de producto" in advertencia
                for advertencia in importacion.advertencias_analisis
            )
        )

    @patch("apps.imports.services.extraer_lista_precios")
    def test_pdf_con_tabla_local_nunca_llama_a_sonnet(self, mock_extraer):
        importacion = self._importar(_pdf_tabular(), "pdf")
        mock_extraer.assert_not_called()
        self.assertFalse(importacion.usa_ia)
        self.assertEqual(importacion.filas.get().codigo, "PDF-1")

    @patch("apps.imports.services.extraer_lista_precios")
    def test_pdf_escaneado_escala_a_sonnet(self, mock_extraer):
        mock_extraer.return_value = (
            [
                {
                    "marca": "Vulcano", "codigo": "ESC-1", "nombre": "Producto escaneado",
                    "descripcion": "", "costo_texto": "2.000,00", "codigo_proveedor": "",
                    "unidad": "", "categoria": "", "nota": "",
                },
            ],
            [],
            {"modelo": "sonnet-test", "input_tokens": 200, "output_tokens": 40},
        )

        importacion = self._importar(_pdf_escaneado(), "pdf")

        mock_extraer.assert_called_once()
        self.assertEqual(mock_extraer.call_args.args[1], "pdf")
        self.assertIsInstance(mock_extraer.call_args.args[0], bytes)
        fila = importacion.filas.get()
        self.assertEqual(fila.confianza, ImportacionFila.Confianza.MEDIA)
        self.assertIn("(IA)", fila.origen)

    @patch("apps.imports.services.extraer_lista_precios")
    def test_docx_sin_tabla_escala_a_sonnet(self, mock_extraer):
        mock_extraer.return_value = (
            [
                {
                    "marca": "", "codigo": "DOC-9", "nombre": "Producto de párrafo",
                    "descripcion": "", "costo_texto": "500", "codigo_proveedor": "",
                    "unidad": "", "categoria": "", "nota": "",
                },
            ],
            [],
            {"modelo": "sonnet-test"},
        )

        importacion = self._importar(_docx_sin_tabla_util(), "docx")

        mock_extraer.assert_called_once()
        self.assertEqual(mock_extraer.call_args.args[1], "docx_texto")
        self.assertIsInstance(mock_extraer.call_args.args[0], str)
        self.assertEqual(importacion.filas.get().codigo, "DOC-9")

    @patch("apps.imports.services.extraer_lista_precios")
    def test_imagen_siempre_llama_a_sonnet_y_guarda_evidencia(self, mock_extraer):
        mock_extraer.return_value = (
            [
                {
                    "marca": "", "codigo": "FOTO-1", "nombre": "Producto de la foto",
                    "descripcion": "", "costo_texto": "750", "codigo_proveedor": "",
                    "unidad": "", "categoria": "", "nota": "",
                },
            ],
            [],
            {"modelo": "sonnet-test"},
        )

        importacion = self._importar(_imagen_jpg(), "imagen")

        mock_extraer.assert_called_once()
        self.assertEqual(mock_extraer.call_args.args[1], "imagen")
        self.assertEqual(importacion.imagenes.count(), 1)
        self.assertEqual(importacion.filas.get().codigo, "FOTO-1")

    @patch("apps.imports.services.extraer_lista_precios")
    def test_falla_de_la_api_reproduce_el_comportamiento_de_falla_existente(self, mock_extraer):
        # docx sin tabla y sin imágenes: si la IA también falla, no queda
        # nada (ni filas ni imágenes) — mismo comportamiento que un archivo
        # no interpretable antes de que existiera este camino.
        mock_extraer.side_effect = ExtraccionIAError("timeout simulado")

        importacion = ImportacionListaPrecios.objects.create(
            proveedor=self.proveedor,
            archivo=_docx_sin_tabla_util(),
            tipo_archivo="docx",
            cargado_por=self.usuario,
        )
        from .parsing import ColumnasNoDetectadas

        with self.assertRaises(ColumnasNoDetectadas):
            procesar_importacion(importacion)

        self.assertFalse(importacion.usa_ia)
        self.assertFalse(importacion.filas.exists())


@override_settings(MEDIA_ROOT=_MEDIA_ROOT, ANTHROPIC_API_KEY="sk-ant-test-clave-secreta-nunca-persistida")
class ApiKeyNuncaPersisteTests(TestCase):
    @patch("apps.imports.services.extraer_lista_precios")
    def test_api_key_no_termina_en_ningun_campo_guardado(self, mock_extraer):
        mock_extraer.return_value = (
            [
                {
                    "marca": "", "codigo": "FOTO-1", "nombre": "Producto",
                    "descripcion": "", "costo_texto": "1", "codigo_proveedor": "",
                    "unidad": "", "categoria": "", "nota": "",
                },
            ],
            [],
            {"modelo": "sonnet-test"},
        )
        usuario = _usuario_admin()
        proveedor = Proveedor.objects.create(nombre_comercial="Proveedor IA", activo=True)
        importacion = ImportacionListaPrecios.objects.create(
            proveedor=proveedor,
            archivo=_imagen_jpg(),
            tipo_archivo="imagen",
            cargado_por=usuario,
        )
        procesar_importacion(importacion)

        from apps.audit.models import AuditLog

        clave = "sk-ant-test-clave-secreta-nunca-persistida"
        self.assertNotIn(clave, str(importacion.ia_resultado))
        for registro in AuditLog.objects.all():
            self.assertNotIn(clave, registro.detalle)


@override_settings(MEDIA_ROOT=_MEDIA_ROOT, ANTHROPIC_API_KEY="sk-ant-test-clave-secreta")
class ArchivoRealListaGeneralTests(TestCase):
    """
    DOC-20260731-WA0188.xlsx es el archivo real (lista de precios de un
    proveedor de climatización, ~400 filas, 30 subtablas tipo "TUBO FUSION
    PN 12" / "TUBO FUSION PN 20" / etc.) que expuso el problema original de
    Excel "tipo catálogo" que motivó tanto el fallback estructural por
    bloques de 11H (_extraer_filas_excel_por_bloques) como, más tarde, el
    camino de IA de esta etapa. No puede quedar sin cobertura de regresión.

    Verificado corriendo analizar_archivo() contra el archivo real antes de
    escribir este test (no de memoria): contra lo que se podía sospechar a
    simple vista, este archivo NO dispara ningún PendienteIA. La fila 18
    (header de la primera subtabla) tiene el TÍTULO DE LA CATEGORÍA
    ("TUBO FUSION PN 12") en la columna de código y la columna de nombre
    vacía, así que detectar_columnas() nunca la reconoce como header
    clásico — ni ella ni ninguna de las otras 29. _contar_filas_header()
    da 0 (no ≥2) para las 400 filas de la hoja, así que el heurístico que
    escala a PendienteIA(tipo="extraccion") cuando hay ≥2 headers clásicos
    independientes nunca se activa. El archivo se resuelve enteramente por
    el fallback estructural por bloques ya existente desde 11H (sin red,
    sin Claude): 149 productos en 30 bloques. Si algún cambio futuro a
    _contar_filas_header/_extraer_filas_excel_por_bloques hiciera que este
    archivo dejara de resolverse localmente y pasara a depender de IA,
    este test lo va a marcar rompiendo el `assert_not_called()`.

    Todas las filas quedan en confianza MEDIA (no ALTA): salen del
    fallback por bloques, que infiere código/nombre por heurística
    posicional en vez de leerlos de un alias exacto — la degradación de
    _extraer_filas_excel_por_bloques() a "media" es incondicional (no
    depende de qué tan claro haya sido el scoring), encontrada al analizar
    Lista_de_Precios_-_Cotizador_OFITT_AGOSTO.xlsm (ver
    ArchivoRealCotizadorOfittAgostoTests). En este archivo no cambia
    ninguna de las categorías/conteos de arriba porque las 149 filas ya
    eran PARA_REVISAR o ERROR (ambas fuerzan incluir=False de por sí,
    independientemente de la confianza).
    """

    def setUp(self):
        self.usuario = _usuario_admin("diego_real")
        self.proveedor = Proveedor.objects.create(nombre_comercial="Proveedor Real", activo=True)

    @patch("apps.imports.services.extraer_lista_precios")
    @patch("apps.imports.services.mapear_columnas")
    def test_lista_general_real_resuelve_local_sin_ia(self, mock_mapear, mock_extraer):
        archivo = SimpleUploadedFile(
            "DOC-20260731-WA0188.xlsx",
            (_FIXTURES_DIR / "DOC-20260731-WA0188.xlsx").read_bytes(),
            content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        )
        importacion = ImportacionListaPrecios.objects.create(
            proveedor=self.proveedor,
            archivo=archivo,
            tipo_archivo="xlsx",
            cargado_por=self.usuario,
        )
        procesar_importacion(importacion)

        mock_mapear.assert_not_called()
        mock_extraer.assert_not_called()
        self.assertFalse(importacion.usa_ia)
        self.assertEqual(importacion.filas.count(), 149)
        self.assertEqual(
            importacion.filas.filter(categoria=ImportacionFila.Categoria.PARA_REVISAR).count(),
            141,
        )
        self.assertEqual(
            importacion.filas.filter(categoria=ImportacionFila.Categoria.ERROR).count(),
            8,
        )
        self.assertTrue(
            all(
                fila.confianza == ImportacionFila.Confianza.MEDIA
                for fila in importacion.filas.all()
            )
        )


@override_settings(MEDIA_ROOT=_MEDIA_ROOT, ANTHROPIC_API_KEY="sk-ant-test-clave-secreta")
class ArchivoRealCotizadorBaseTests(TestCase):
    """
    Cotizador_Base.xlsm es el archivo .xlsm real (Importación #5) que
    expuso el pre-filtro de _fila_parece_producto_confiable: tiene 4
    hojas — "Cotizador" (header único en fila 6: Rubro, SubRubro, Código,
    Descripción, Código EAN, Precio Unit., Cantidad, Total Neto),
    "Condiciones" (texto de términos y condiciones, sin ningún dato de
    producto) y "MARA" (un extracto tipo SAP de ~9600 filas, sin ninguna
    columna de precio, pero con ID/EAN/peso numéricos que antes colaban
    como "costo" para parsear_costo()). "GRILLA CLIENTE" resuelve local
    también (marca no detectada, categoría PARA_REVISAR, sin pasar por IA).

    Antes del pre-filtro, "Condiciones" y "MARA" escalaban a Sonnet
    (confirmado contra el ia_resultado real de la Importación #5: 1.493 +
    184.326 = 185.819 tokens de entrada, ~USD 0.37, solo para que Claude
    concluyera "esto no tiene precios"). Mismo criterio de regresión
    permanente que ArchivoRealListaGeneralTests para DOC-20260731-WA0188:
    si algún cambio futuro a _fila_parece_producto_confiable o a los
    caminos locales hiciera que este archivo dejara de resolverse sin
    IA, este test lo va a marcar rompiendo los `assert_not_called()`.

    También es el fixture real que expuso la detección de moneda: la
    columna "Precio Unit." de "Cotizador" tiene number_format de USD
    (`"U$S"\ #,##0.00;...`) en Excel, invisible como texto.

    Este mismo fixture es, sin haberlo buscado, el que expuso el bug de
    ALIAS_COSTO que no reconocía "Precio Unit." (encontrado analizando
    Lista_de_Precios_-_Cotizador_OFITT_AGOSTO.xlsm, ver
    ArchivoRealCotizadorOfittAgostoTests más abajo): antes del fix,
    "Cotizador" caía al fallback por bloques y terminaba usando "Rubro"
    como nombre y "Código EAN" como código para sus 1061 filas — nunca
    se había notado acá porque este test solo verificaba cantidades y
    moneda, no el valor de código/nombre. Con el fix, "Cotizador" resuelve
    por el camino clásico (código→"Código", nombre→"Descripción",
    categoria→"Rubro", todos correctos) y recupera además 17 filas que el
    fallback por bloques descartaba por no pasar su propio filtro de
    "parece código"/"parece nombre" (1078 en vez de 1061).
    """

    def setUp(self):
        self.usuario = _usuario_admin("diego_cotizador")
        self.proveedor = Proveedor.objects.create(nombre_comercial="Proveedor Cotizador", activo=True)

    @patch("apps.imports.services.extraer_lista_precios")
    @patch("apps.imports.services.mapear_columnas")
    def test_cotizador_base_resuelve_cotizador_local_y_filtra_mara_y_condiciones(
        self, mock_mapear, mock_extraer
    ):
        archivo = SimpleUploadedFile(
            "Cotizador_Base.xlsm",
            (_FIXTURES_DIR / "Cotizador_Base.xlsm").read_bytes(),
            content_type="application/vnd.ms-excel.sheet.macroEnabled.12",
        )
        importacion = ImportacionListaPrecios.objects.create(
            proveedor=self.proveedor,
            archivo=archivo,
            tipo_archivo="xlsm",
            cargado_por=self.usuario,
        )
        procesar_importacion(importacion)

        mock_mapear.assert_not_called()
        mock_extraer.assert_not_called()
        self.assertFalse(importacion.usa_ia)

        advertencias = importacion.advertencias_analisis
        self.assertTrue(
            any(
                adv.startswith("Hoja MARA:")
                and "no se detectó ninguna fila con estructura de producto" in adv
                for adv in advertencias
            ),
            advertencias,
        )
        self.assertTrue(
            any(
                adv.startswith("Hoja Condiciones:")
                and "no se detectó ninguna fila con estructura de producto" in adv
                for adv in advertencias
            ),
            advertencias,
        )

        # "Cotizador" (1078, camino clásico desde el fix de ALIAS_COSTO) +
        # "GRILLA CLIENTE" (1078, clásico sin marca) resuelven local, sin
        # depender de la IA para nada. "MARA" y "Condiciones" no aportan
        # ninguna fila (arriba).
        self.assertEqual(importacion.filas.count(), 2156)
        self.assertEqual(
            importacion.filas.filter(categoria=ImportacionFila.Categoria.PARA_REVISAR).count(),
            1046,
        )
        self.assertEqual(
            importacion.filas.filter(categoria=ImportacionFila.Categoria.ERROR).count(),
            11,
        )

        # Detección de moneda vía number_format: las 1078 filas de
        # "Cotizador" están en USD en la planilla real (celda por celda,
        # no una muestra) y quedan en confianza ALTA — USD confiable no
        # baja confianza ni fuerza revisión manual.
        filas_cotizador = importacion.filas.filter(origen__startswith="Hoja Cotizador")
        self.assertEqual(filas_cotizador.count(), 1078)
        self.assertEqual(
            filas_cotizador.filter(moneda=Moneda.USD, confianza=ImportacionFila.Confianza.ALTA).count(),
            1078,
        )

        # Código y nombre reales (no "Rubro" ni "Código EAN" del bug que
        # este mismo archivo tenía antes del fix de ALIAS_COSTO).
        primera_cotizador = filas_cotizador.order_by("numero_fila").first()
        self.assertEqual(primera_cotizador.codigo, "10000005")
        self.assertEqual(primera_cotizador.nombre_texto, 'VALVULA 1/2""x16  REC MANUAL"')
        self.assertEqual(primera_cotizador.categoria_texto, "Accesorio")
        # "GRILLA CLIENTE" no tiene señal de moneda en el archivo real:
        # queda en el default ARS.
        self.assertEqual(
            importacion.filas.exclude(origen__startswith="Hoja Cotizador")
            .exclude(moneda=Moneda.ARS)
            .count(),
            0,
        )


@override_settings(MEDIA_ROOT=_MEDIA_ROOT, ANTHROPIC_API_KEY="sk-ant-test-clave-secreta")
class ArchivoRealCotizadorOfittAgostoTests(TestCase):
    """
    Cotizador_OFITT_Agosto.xlsm es el archivo real que expuso el bug de
    ALIAS_COSTO: su hoja "Cotizador" (header único en fila 6) usa
    literalmente "Código", "Descripción" y "Precio Unit." — Código y
    Descripción ya estaban cubiertos por ALIAS_CODIGO/ALIAS_NOMBRE, pero
    "Precio Unit." (abreviado) no estaba en ALIAS_COSTO, que solo tenía
    la forma larga "precio unitario". Esa única columna faltante hacía
    fallar detectar_columnas() para toda la fila 6 (exige código+nombre+
    costo juntos) y la hoja caía al fallback estructural por bloques
    (_extraer_filas_excel_por_bloques / _inferir_columnas_bloque), que
    infiere columnas por heurística posicional en vez de por alias: acá
    elegía "SubRubro" como nombre (empate 188/188 contra "Descripción",
    resuelto por orden de columna) y "Código EAN" como código (puntaje
    188 contra 164 de la columna "Código" real, porque un EAN de 13
    dígitos matchea el patrón de "parece código" con más consistencia
    que un código alfanumérico como "MTPEX1620BL-R12"). Confirmado
    ejecutando el parser real contra este archivo antes de escribir este
    test, no de memoria.

    Con el fix (agregar "precio unit." a ALIAS_COSTO), la hoja resuelve
    por el camino clásico: código→"Código" (col. F), nombre→"Descripción"
    (col. G) — nunca "Código EAN" (col. H) ni "SubRubro" (col. A).
    """

    def setUp(self):
        self.usuario = _usuario_admin("diego_ofitt")
        self.proveedor = Proveedor.objects.create(nombre_comercial="Proveedor OFITT", activo=True)

    @patch("apps.imports.services.extraer_lista_precios")
    @patch("apps.imports.services.mapear_columnas")
    def test_cotizador_ofitt_agosto_resuelve_clasico_con_codigo_y_nombre_reales(
        self, mock_mapear, mock_extraer
    ):
        archivo = SimpleUploadedFile(
            "Cotizador_OFITT_Agosto.xlsm",
            (_FIXTURES_DIR / "Cotizador_OFITT_Agosto.xlsm").read_bytes(),
            content_type="application/vnd.ms-excel.sheet.macroEnabled.12",
        )
        importacion = ImportacionListaPrecios.objects.create(
            proveedor=self.proveedor,
            archivo=archivo,
            tipo_archivo="xlsm",
            cargado_por=self.usuario,
        )
        procesar_importacion(importacion)

        mock_mapear.assert_not_called()
        mock_extraer.assert_not_called()
        self.assertFalse(importacion.usa_ia)

        advertencias = importacion.advertencias_analisis
        self.assertFalse(
            any("bloque(s) comerciales" in adv for adv in advertencias),
            advertencias,
        )

        filas_cotizador = importacion.filas.filter(origen__startswith="Hoja Cotizador").order_by(
            "numero_fila"
        )
        self.assertEqual(filas_cotizador.count(), 188)

        primera = filas_cotizador.first()
        self.assertEqual(primera.codigo, "MTPEX1620BL-R12")
        self.assertEqual(
            primera.nombre_texto, "TUBO PE-X 5 CAPAS Ø16X2,0 - 120M - BLANC"
        )

        ultima = filas_cotizador.last()
        self.assertEqual(ultima.codigo, "VALV.ZONA")
        self.assertEqual(ultima.nombre_texto, "VÁLVULA DE ZONA DE 3 VÍAS UNIVERSAL")

        # Ningún código real quedó tomado de la columna "Código EAN"
        # (siempre 13 dígitos numéricos) ni ningún nombre quedó tomado de
        # "SubRubro" (categorías cortas repetidas como "PE-X").
        self.assertFalse(
            any(len(f.codigo) == 13 and f.codigo.isdigit() for f in filas_cotizador)
        )


@override_settings(ANTHROPIC_API_KEY="sk-ant-test-clave-secreta")
class LlamarToolForzadoErroresTests(TestCase):
    """
    A diferencia del resto de este archivo (que mockea services.mapear_columnas/
    services.extraer_lista_precios, es decir, la frontera de apps.imports.ai
    completa), estos dos tests ejercitan apps.imports.ai._llamar_tool_forzado()
    de verdad, mockeando solo el cliente de anthropic — porque lo que se
    prueba es comportamiento interno de esa función (traducción de
    APITimeoutError, rechazo de una respuesta cortada por max_tokens), no
    algo que se pueda observar mockeando un nivel más arriba.
    """

    def _mockear_cliente(self, mock_anthropic_cls, *, side_effect=None, respuesta=None):
        cliente = MagicMock()
        cliente.with_options.return_value = cliente
        if side_effect is not None:
            cliente.messages.create.side_effect = side_effect
        else:
            cliente.messages.create.return_value = respuesta
        mock_anthropic_cls.return_value = cliente
        return cliente

    @patch("apps.imports.ai.anthropic.Anthropic")
    def test_timeout_de_la_api_se_traduce_a_extraccioniaerror_especifico(self, mock_anthropic_cls):
        # APITimeoutError hereda de APIConnectionError (verificado contra el
        # anthropic==1.2.0 instalado): esto confirma que el isinstance()
        # dentro del except existente distingue el mensaje sin necesitar un
        # except propio.
        self._mockear_cliente(
            mock_anthropic_cls,
            side_effect=anthropic.APITimeoutError(request=MagicMock()),
        )

        with self.assertRaisesMessage(ExtraccionIAError, "tardó demasiado"):
            mapear_columnas(
                ["Marca", "Código", "Nombre", "Costo"],
                [["Vulcano", "1", "Producto", "10"]],
            )

    @patch("apps.imports.ai.anthropic.Anthropic")
    def test_respuesta_cortada_por_max_tokens_no_se_usa_como_extraccion_parcial(
        self, mock_anthropic_cls
    ):
        respuesta = MagicMock()
        respuesta.stop_reason = "max_tokens"
        self._mockear_cliente(mock_anthropic_cls, respuesta=respuesta)

        with self.assertRaisesMessage(ExtraccionIAError, "límite de tokens"):
            mapear_columnas(
                ["Marca", "Código", "Nombre", "Costo"],
                [["Vulcano", "1", "Producto", "10"]],
            )


class ExtraerListaPreciosRedDeContencionTests(TestCase):
    """
    Caso real: Importación #10 (Uriarte Taldea, catálogo en grilla de
    tarjetas) — Sonnet devolvió productos:[] Y advertencias:[] a la vez,
    sin explicar por qué. Confirmado ejecutando extraer_lista_precios()
    contra el PDF real con la API real (no de memoria): input_tokens
    coincidía casi exacto con el ia_resultado guardado en la Importación
    #10 (26915), y la respuesta también vino con los dos arrays vacíos.

    El ajuste de _SISTEMA_EXTRACCION (grillas de tarjetas + instrucción de
    siempre explicarse) apunta a que esto no vuelva a pasar, pero esta red
    de contención en extraer_lista_precios() no depende de que el prompt
    funcione: mockea el cliente de anthropic directamente (mismo patrón
    que LlamarToolForzadoErroresTests) para forzar el caso "los dos arrays
    vacíos" sin pegarle a la API real, y confirma que nunca llega vacío de
    contexto al resto del pipeline.
    """

    def _mockear_respuesta_vacia(self, mock_anthropic_cls):
        bloque = MagicMock()
        bloque.type = "tool_use"
        bloque.name = "extraer_lista_precios"
        bloque.input = {"productos": [], "advertencias": []}

        respuesta = MagicMock()
        respuesta.stop_reason = "tool_use"
        respuesta.content = [bloque]
        respuesta.model = "claude-sonnet-5"
        respuesta.usage.input_tokens = 26915
        respuesta.usage.output_tokens = 59

        cliente = MagicMock()
        cliente.with_options.return_value = cliente
        cliente.messages.create.return_value = respuesta
        mock_anthropic_cls.return_value = cliente

    @patch("apps.imports.ai.anthropic.Anthropic")
    def test_productos_y_advertencias_vacios_sintetiza_advertencia_generica(
        self, mock_anthropic_cls
    ):
        self._mockear_respuesta_vacia(mock_anthropic_cls)

        productos, advertencias, meta = extraer_lista_precios(b"%PDF-1.4 ...", "pdf")

        self.assertEqual(productos, [])
        self.assertEqual(len(advertencias), 1)
        self.assertIn("no extrajo productos", advertencias[0])
        self.assertIn("revisar manualmente", advertencias[0])
        self.assertEqual(meta["input_tokens"], 26915)

    @patch("apps.imports.ai.anthropic.Anthropic")
    def test_no_pisa_una_advertencia_real_del_modelo(self, mock_anthropic_cls):
        bloque = MagicMock()
        bloque.type = "tool_use"
        bloque.name = "extraer_lista_precios"
        bloque.input = {
            "productos": [],
            "advertencias": ["El documento es una grilla de tarjetas ilegible."],
        }
        respuesta = MagicMock()
        respuesta.stop_reason = "tool_use"
        respuesta.content = [bloque]
        respuesta.model = "claude-sonnet-5"
        respuesta.usage.input_tokens = 100
        respuesta.usage.output_tokens = 20
        cliente = MagicMock()
        cliente.with_options.return_value = cliente
        cliente.messages.create.return_value = respuesta
        mock_anthropic_cls.return_value = cliente

        productos, advertencias, meta = extraer_lista_precios(b"%PDF-1.4 ...", "pdf")

        self.assertEqual(advertencias, ["El documento es una grilla de tarjetas ilegible."])
