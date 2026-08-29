import csv
import hashlib
import re
import unicodedata
import warnings
import zipfile
from dataclasses import dataclass, field
from decimal import Decimal, InvalidOperation
from io import BytesIO
from pathlib import Path

import openpyxl
import xlrd
from docx import Document
from openpyxl.utils import get_column_letter
from PIL import Image, ImageOps
from pypdf import PdfReader


MAX_ARCHIVO_BYTES = 30 * 1024 * 1024
MAX_FILAS = 10_000
MAX_PAGINAS_PDF = 250
MAX_IMAGENES = 60
MAX_IMAGEN_BYTES = 8 * 1024 * 1024
MAX_IMAGEN_PIXELES = 24_000_000
MAX_ZIP_DESCOMPRIMIDO = 180 * 1024 * 1024
MAX_RATIO_ZIP = 250
MAX_CARACTERES_TEXTO_IA = 300_000

ALIAS_MARCA = {"marca", "brand", "fabricante", "manufacturer"}
ALIAS_CODIGO = {
    "codigo", "código", "code", "sku", "modelo", "model", "referencia", "ref",
    "cod", "cod.", "cod. fabricante", "cod fabricante", "codigo fabricante",
    "código fabricante", "part number", "part no", "pn",
}
ALIAS_NOMBRE = {
    "nombre", "descripcion", "descripción", "producto", "detalle", "articulo",
    "artículo", "item", "ítem", "concepto",
}
ALIAS_COSTO = {
    "costo", "precio", "precio neto", "precio unitario", "importe",
    "precio de costo", "neto", "precio lista", "precio de lista", "price",
    "unit price", "p. unitario", "precio bonif", "precio bonificado",
    "precio final", "costo neto", "costo bonificado",
}
ALIAS_CODIGO_PROVEEDOR = {
    "codigo proveedor", "código proveedor", "cod. proveedor", "cod proveedor",
    "sku proveedor", "codigo interno", "código interno", "cod interno",
}
ALIAS_UNIDAD = {
    "unidad", "unidad de medida", "u.m.", "um", "unit", "medida",
}
ALIAS_CATEGORIA = {
    "categoria", "categoría", "familia", "rubro", "linea", "línea", "grupo",
}
ALIAS_DESCRIPCION_EXTENDIDA = {
    "descripcion extendida", "descripción extendida", "descripcion tecnica",
    "descripción técnica", "caracteristicas", "características", "observaciones",
}

CAMPOS_REQUERIDOS = ("codigo", "nombre", "costo")
CAMPOS_ALIAS = {
    "marca": ALIAS_MARCA,
    "codigo": ALIAS_CODIGO,
    "nombre": ALIAS_NOMBRE,
    "costo": ALIAS_COSTO,
    "codigo_proveedor": ALIAS_CODIGO_PROVEEDOR,
    "unidad": ALIAS_UNIDAD,
    "categoria": ALIAS_CATEGORIA,
    "descripcion": ALIAS_DESCRIPCION_EXTENDIDA,
}


@dataclass
class FilaAnalizada:
    numero: int
    origen: str
    datos: dict
    confianza: str = "alta"
    nota_ia: str = ""


@dataclass
class ImagenAnalizada:
    contenido: bytes
    extension: str
    origen: str
    nombre_original: str = ""
    numero_fila_origen: int | None = None
    ancho: int | None = None
    alto: int | None = None
    huella_sha256: str = ""


@dataclass
class PendienteIA:
    """
    Trabajo que este módulo no pudo resolver de forma local y que requiere
    llamar a Claude. parsing.py arma el pendiente pero nunca lo resuelve —
    eso es responsabilidad de apps.imports.services (única capa que tiene
    acceso a DB para el cache de ProveedorColumnMapping) llamando a
    apps.imports.ai (única capa que habla con la API).

    tipo="mapeo_columnas": Excel/CSV "plano" cuyos encabezados no matchean
    los alias conocidos. `matriz_restante` son las filas después del header
    candidato, listas para pasarle a extraer_filas_con_mapeo() una vez que
    se resuelva el mapeo.

    tipo="extraccion": PDF/imagen/docx/Excel-catálogo. `contenido` son bytes
    (pdf/imagen) o texto ya extraído localmente (docx_texto/excel_celdas) —
    nunca el binario original para estos dos últimos.
    """

    tipo: str
    origen: str
    encabezados: list | None = None
    muestra_filas: list | None = None
    matriz_restante: list | None = None
    contenido: object = None
    tipo_contenido: str | None = None
    extension_imagen: str | None = None


@dataclass
class ResultadoAnalisis:
    filas: list[FilaAnalizada] = field(default_factory=list)
    imagenes: list[ImagenAnalizada] = field(default_factory=list)
    advertencias: list[str] = field(default_factory=list)
    pendientes_ia: list[PendienteIA] = field(default_factory=list)


def _normalizar(texto):
    texto = str(texto or "").strip().lower()
    texto = "".join(
        c
        for c in unicodedata.normalize("NFD", texto)
        if unicodedata.category(c) != "Mn"
    )
    texto = re.sub(r"\s+", " ", texto)
    return texto


_ALIAS_NORMALIZADOS = {
    campo: {_normalizar(alias) for alias in aliases}
    for campo, aliases in CAMPOS_ALIAS.items()
}


class ColumnasNoDetectadas(Exception):
    def __init__(self, faltantes, encabezados, origen=""):
        self.faltantes = faltantes
        self.encabezados = [str(h) for h in encabezados if str(h or "").strip()]
        prefijo = f"{origen}: " if origen else ""
        mensaje = (
            f"{prefijo}No se pudieron reconocer las columnas: {', '.join(faltantes)}. "
            f"Encabezados encontrados: {', '.join(self.encabezados) or '(ninguno)'}."
        )
        super().__init__(mensaje)


class ArchivoImportacionInvalido(ValueError):
    pass


EXTENSIONES_IMAGEN = {"jpg", "jpeg", "png", "webp"}


def tipo_archivo_por_nombre(nombre):
    extension = Path(nombre or "").suffix.lower().lstrip(".")
    if extension in EXTENSIONES_IMAGEN:
        return "imagen"
    if extension not in {"xlsx", "xls", "csv", "pdf", "docx"}:
        raise ArchivoImportacionInvalido(
            "Formato no soportado. Usá .xlsx, .xls, .csv, .pdf, .docx, .jpg, .jpeg, .png o .webp."
        )
    return extension


def detectar_columnas(encabezados):
    """
    Detecta columnas por aliases. Código, nombre y costo son obligatorios.

    Marca dejó de ser obligatoria en 11H: algunos proveedores identifican
    inequívocamente por su propio código y otros mandan listas de una sola
    marca. Si falta, la fila se conserva para revisión en vez de descartar
    todo el archivo.
    """
    normalizados = [_normalizar(h) for h in encabezados]
    mapeo = {}
    faltantes = []

    for campo, aliases in _ALIAS_NORMALIZADOS.items():
        indice = next(
            (i for i, valor in enumerate(normalizados) if valor in aliases),
            None,
        )
        if indice is not None:
            mapeo[campo] = indice
        elif campo in CAMPOS_REQUERIDOS:
            faltantes.append(campo)

    if faltantes:
        raise ColumnasNoDetectadas(faltantes, encabezados)
    return mapeo


def _puntaje_encabezado(encabezados):
    normalizados = [_normalizar(h) for h in encabezados]
    return sum(
        1
        for aliases in _ALIAS_NORMALIZADOS.values()
        if any(valor in aliases for valor in normalizados)
    )


def _mejor_candidata_encabezado(matriz, limite=25):
    """
    Fila con más alias de campos canónicos reconocidos, sin exigir que estén
    los 3 obligatorios juntos (a diferencia de detectar_columnas). Se usa
    tanto para el mensaje de ColumnasNoDetectadas como para decidir, cuando
    el mapeo clásico falla, qué fila mandarle a Claude como candidata a
    encabezado y desde dónde empezar a leer datos.
    """
    mejor_puntaje = 0
    mejor_posicion = 0
    mejor_fila = []
    for posicion, (_, fila) in enumerate(matriz[:limite]):
        puntaje = _puntaje_encabezado(fila)
        if puntaje > mejor_puntaje:
            mejor_puntaje = puntaje
            mejor_posicion = posicion
            mejor_fila = fila
    return mejor_posicion, mejor_fila


def _detectar_encabezado(matriz, origen, limite=25):
    for posicion, (_, fila) in enumerate(matriz[:limite]):
        try:
            return posicion, detectar_columnas(fila)
        except ColumnasNoDetectadas:
            continue
    _, mejor_fila = _mejor_candidata_encabezado(matriz, limite=limite)
    raise ColumnasNoDetectadas(list(CAMPOS_REQUERIDOS), mejor_fila, origen=origen)


def parsear_costo(valor):
    """
    Tolera números, separadores argentinos/internacionales y texto con
    símbolos de moneda. Devuelve None si no se puede interpretar.
    """
    if valor is None or valor == "":
        return None

    if isinstance(valor, (int, float, Decimal)):
        try:
            return Decimal(str(valor)).quantize(Decimal("0.01"))
        except (InvalidOperation, ValueError):
            return None

    texto = str(valor).replace("\xa0", " ").strip()
    if not texto:
        return None

    coincidencia = re.search(r"-?\(?\d[\d.,\s]*\)?", texto)
    if not coincidencia:
        return None

    numero = coincidencia.group(0).strip().replace(" ", "")
    negativo_parentesis = numero.startswith("(") and numero.endswith(")")
    numero = numero.strip("()")

    if "," in numero and "." in numero:
        if numero.rfind(",") > numero.rfind("."):
            numero = numero.replace(".", "").replace(",", ".")
        else:
            numero = numero.replace(",", "")
    elif "," in numero:
        numero = numero.replace(".", "").replace(",", ".")

    try:
        valor_decimal = Decimal(numero)
        if negativo_parentesis:
            valor_decimal = -valor_decimal
        return valor_decimal.quantize(Decimal("0.01"))
    except (InvalidOperation, ValueError):
        return None


def _texto_celda(valor):
    if valor is None:
        return ""
    if isinstance(valor, float) and valor.is_integer():
        return str(int(valor))
    return str(valor).strip()


def extraer_filas_con_mapeo(matriz_restante, origen, mapeo, confianza="alta"):
    """
    Lee filas de datos dado un mapeo campo→índice de columna ya resuelto —
    por detectar_columnas() (camino clásico) o por ai.mapear_columnas()
    (camino Haiku, cuando los alias conocidos no reconocen el encabezado).

    Es la ÚNICA función que interpreta filas de producto a partir de un
    mapeo de columnas; tanto _extraer_filas_matriz como el resolutor de
    PendienteIA en services.py pasan por acá para no duplicar esta lógica.
    """
    advertencias = []
    filas = []
    for numero, valores in matriz_restante:
        if len(filas) >= MAX_FILAS:
            advertencias.append(
                f"{origen}: se alcanzó el límite de {MAX_FILAS} filas; el resto no se analizó."
            )
            break
        if not valores or all(not _texto_celda(v) for v in valores):
            continue

        # Evita repetir un encabezado cuando una lista imprime cabecera en cada página.
        try:
            detectar_columnas(valores)
        except ColumnasNoDetectadas:
            pass
        else:
            continue

        def obtener(campo):
            indice = mapeo.get(campo)
            if indice is None or indice >= len(valores):
                return ""
            return valores[indice]

        costo_crudo = obtener("costo")
        filas.append(
            FilaAnalizada(
                numero=max(1, int(numero)),
                origen=origen,
                confianza=confianza,
                datos={
                    "marca": _texto_celda(obtener("marca")),
                    "codigo": _texto_celda(obtener("codigo")),
                    "nombre": _texto_celda(obtener("nombre")),
                    "descripcion": _texto_celda(obtener("descripcion")),
                    "costo_crudo": costo_crudo,
                    "codigo_proveedor": _texto_celda(obtener("codigo_proveedor")),
                    "unidad": _texto_celda(obtener("unidad")),
                    "categoria": _texto_celda(obtener("categoria")),
                },
            )
        )
    return filas, advertencias


def _extraer_filas_matriz(matriz, origen, confianza="alta"):
    if not matriz:
        return [], []

    posicion_header, mapeo = _detectar_encabezado(matriz, origen)
    advertencias = []
    if "marca" not in mapeo:
        advertencias.append(
            f"{origen}: no se detectó columna Marca. Las filas nuevas deberán "
            "tener una marca resoluble antes de confirmarse."
        )

    filas, adv_filas = extraer_filas_con_mapeo(
        matriz[posicion_header + 1 :], origen, mapeo, confianza=confianza
    )
    advertencias.extend(adv_filas)
    return filas, advertencias




PRECIO_PRIORITARIO = {
    "precio bonif",
    "precio bonificado",
    "precio neto",
    "precio final",
    "costo neto",
    "costo bonificado",
}

ENCABEZADOS_NO_CATEGORIA = {
    "embalaje",
    "pedido",
    "sub total",
    "subtotal",
    "cantidad",
    "cant",
}


def _indice_precio_bloque(valores):
    """
    Detecta la columna de precio/costo de una cabecera de bloque.

    Si la lista tiene simultáneamente precio de lista y precio bonificado,
    prioriza el bonificado/neto porque representa mejor el costo efectivo
    que ARQCLIMA pagaría al proveedor.
    """
    candidatos = []
    for indice, valor in enumerate(valores):
        normalizado = _normalizar(valor)
        if not normalizado:
            continue

        if normalizado in PRECIO_PRIORITARIO:
            candidatos.append((0, indice))
            continue

        if normalizado in _ALIAS_NORMALIZADOS["costo"]:
            candidatos.append((1, indice))
            continue

        if normalizado.startswith("precio") or normalizado.startswith("costo"):
            candidatos.append((2, indice))

    if not candidatos:
        return None
    candidatos.sort()
    return candidatos[0][1]


def _parece_codigo(valor):
    texto = _texto_celda(valor)
    if not texto or len(texto) < 5:
        return False
    if len(texto) > 80:
        return False
    if texto.count(" ") > 2:
        return False

    letras = sum(1 for caracter in texto if caracter.isalpha())
    digitos = sum(1 for caracter in texto if caracter.isdigit())
    if digitos >= 5 and letras <= 12:
        return True

    return bool(
        re.fullmatch(r"[A-Za-z0-9][A-Za-z0-9._/-]{4,}", texto)
        and digitos >= 1
    )


def _parece_nombre_producto(valor):
    texto = _texto_celda(valor)
    if len(texto) < 4:
        return False
    return any(caracter.isalpha() for caracter in texto)


def _fila_parece_producto(valores):
    """
    Señal barata e independiente del encabezado: ¿esta fila tiene "algo que
    parece código", "algo que parece nombre" y "algo que parece un precio"?
    No identifica QUÉ columna es cada cosa (para eso está _inferir_columnas_bloque),
    solo estima cuántas filas de una hoja parecen filas de producto, para
    medir si el camino local (clásico + bloques) cubrió razonablemente el
    archivo o si hace falta escalar a IA.
    """
    return (
        any(_parece_codigo(v) for v in valores)
        and any(_parece_nombre_producto(v) for v in valores)
        and any(parsear_costo(v) is not None for v in valores)
    )


def _contar_filas_header(matriz, tope=2):
    """
    Cuenta filas de TODA la matriz (no solo las primeras 25) donde
    detectar_columnas() tiene éxito — es decir, headers independientes que
    matchean código+nombre+costo por sí solos. Una tabla plana tiene 1; un
    Excel "tipo catálogo" con una subtabla por sección (ej. "TUBO FUSION
    PN12", "TUBO FUSION PN20", "CODO A 90º"... cada una con su propio
    header) tiene varios, y ningún mapeo de columnas único cubre el archivo.
    Corta apenas llega a `tope` porque solo nos importa distinguir 0/1 de 2+.
    """
    contador = 0
    for _, valores in matriz:
        try:
            detectar_columnas(valores)
        except ColumnasNoDetectadas:
            continue
        contador += 1
        if contador >= tope:
            break
    return contador


def _inferir_columnas_bloque(filas_bloque, indice_precio):
    """
    Infere Código y Nombre a partir del patrón de las filas del bloque.

    Sirve para listas comerciales donde el proveedor no rotula esas dos
    columnas en cada cabecera, pero las repite de manera estructurada.
    """
    candidatas = [
        (numero, valores)
        for numero, valores in filas_bloque
        if indice_precio < len(valores)
        and parsear_costo(valores[indice_precio]) is not None
    ]
    if not candidatas:
        return None, None

    limites = range(0, indice_precio)
    puntajes_codigo = {}
    puntajes_nombre = {}

    for indice in limites:
        puntajes_codigo[indice] = sum(
            1
            for _, valores in candidatas
            if indice < len(valores) and _parece_codigo(valores[indice])
        )
        puntajes_nombre[indice] = sum(
            1
            for _, valores in candidatas
            if indice < len(valores) and _parece_nombre_producto(valores[indice])
        )

    indice_codigo = max(puntajes_codigo, key=puntajes_codigo.get)
    if puntajes_codigo[indice_codigo] == 0:
        return None, None

    candidatos_nombre = {
        indice: puntaje
        for indice, puntaje in puntajes_nombre.items()
        if indice != indice_codigo
    }
    if not candidatos_nombre:
        return None, None

    indice_nombre = max(candidatos_nombre, key=candidatos_nombre.get)
    if candidatos_nombre[indice_nombre] == 0:
        return None, None

    return indice_codigo, indice_nombre


def _categoria_bloque(fila_cabecera, indice_precio):
    candidatos = []
    for indice, valor in enumerate(fila_cabecera[:indice_precio]):
        texto = _texto_celda(valor)
        normalizado = _normalizar(texto)
        if not texto or normalizado in ENCABEZADOS_NO_CATEGORIA:
            continue
        if normalizado in _ALIAS_NORMALIZADOS["costo"]:
            continue
        if _parece_codigo(texto):
            continue
        if any(caracter.isalpha() for caracter in texto):
            candidatos.append(texto)

    if not candidatos:
        return ""
    return max(candidatos, key=len)


def _extraer_filas_excel_por_bloques(matriz, origen):
    """
    Fallback estructural para listas Excel armadas por secciones.

    Caso típico:
        [categoría]  EMBALAJE  PRECIO $  PRECIO BONIF  PEDIDO  SUB TOTAL
        [código]     [nombre]  [pack]    [precio]       ...

    En estas planillas Código y Nombre no aparecen como encabezados
    explícitos, por lo que el detector tabular clásico no puede encontrarlos.
    """
    cabeceras = []
    for posicion, (_, valores) in enumerate(matriz):
        indice_precio = _indice_precio_bloque(valores)
        if indice_precio is not None:
            cabeceras.append((posicion, indice_precio))

    if not cabeceras:
        return [], []

    filas_resultado = []
    advertencias = []
    bloques_validos = 0

    for bloque_indice, (posicion, indice_precio) in enumerate(cabeceras):
        fin = (
            cabeceras[bloque_indice + 1][0]
            if bloque_indice + 1 < len(cabeceras)
            else len(matriz)
        )
        numero_cabecera, fila_cabecera = matriz[posicion]
        filas_bloque = matriz[posicion + 1 : fin]
        indice_codigo, indice_nombre = _inferir_columnas_bloque(
            filas_bloque,
            indice_precio,
        )
        if indice_codigo is None or indice_nombre is None:
            continue

        categoria = _categoria_bloque(fila_cabecera, indice_precio)
        indices_precio_fallback = [indice_precio]
        for indice, valor in enumerate(fila_cabecera):
            normalizado = _normalizar(valor)
            if indice == indice_precio:
                continue
            if (
                normalizado in _ALIAS_NORMALIZADOS["costo"]
                or normalizado.startswith("precio")
                or normalizado.startswith("costo")
            ):
                indices_precio_fallback.append(indice)

        bloques_validos += 1

        for numero, valores in filas_bloque:
            if len(filas_resultado) >= MAX_FILAS:
                advertencias.append(
                    f"{origen}: se alcanzó el límite de {MAX_FILAS} filas."
                )
                break

            if (
                indice_codigo >= len(valores)
                or indice_nombre >= len(valores)
                or indice_precio >= len(valores)
            ):
                continue

            codigo = _texto_celda(valores[indice_codigo])
            nombre = _texto_celda(valores[indice_nombre])
            costo_crudo = None
            for indice_costo in indices_precio_fallback:
                if indice_costo >= len(valores):
                    continue
                candidato = valores[indice_costo]
                if parsear_costo(candidato) is not None:
                    costo_crudo = candidato
                    break

            if not _parece_codigo(codigo) or not _parece_nombre_producto(nombre):
                continue
            if costo_crudo is None:
                continue

            filas_resultado.append(
                FilaAnalizada(
                    numero=max(1, int(numero)),
                    origen=f"{origen} · {categoria}" if categoria else origen,
                    confianza="alta",
                    datos={
                        "marca": "",
                        "codigo": codigo,
                        "nombre": nombre,
                        "descripcion": "",
                        "costo_crudo": costo_crudo,
                        "codigo_proveedor": "",
                        "unidad": "",
                        "categoria": categoria,
                    },
                )
            )

    if filas_resultado:
        advertencias.append(
            f"{origen}: se detectaron {len(filas_resultado)} productos en "
            f"{bloques_validos} bloque(s) comerciales aunque la planilla no "
            "trae encabezados explícitos Código/Nombre."
        )
        advertencias.append(
            f"{origen}: cuando existe una columna de precio bonificado/neto "
            "se usa como costo efectivo en lugar del precio de lista."
        )

    return filas_resultado, advertencias


def _recortar_texto_ia(texto, origen):
    if len(texto) <= MAX_CARACTERES_TEXTO_IA:
        return texto, None
    return (
        texto[:MAX_CARACTERES_TEXTO_IA],
        f"{origen}: el contenido para IA se recortó a {MAX_CARACTERES_TEXTO_IA} caracteres.",
    )


def construir_texto_celdas_no_vacias(matriz, origen=""):
    """
    Representación de solo celdas no vacías (coordenada Excel: valor), nunca
    el binario del archivo. Para un Excel "tipo catálogo" real de ~400 filas
    x 29 columnas con ~150 productos esto da unos pocos miles de tokens, muy
    lejos del costo de mandar el archivo completo.

    Se preservan los cortes de fila en blanco como línea vacía: son la
    misma señal visual que ve una persona para separar bloques/secciones,
    y ayudan al modelo a agrupar filas del mismo producto.
    """
    lineas = []
    hubo_fila_no_vacia = False
    fila_en_blanco_pendiente = False
    for numero, valores in matriz:
        celdas = [
            (indice, texto)
            for indice, valor in enumerate(valores)
            if (texto := _texto_celda(valor))
        ]
        if not celdas:
            if hubo_fila_no_vacia:
                fila_en_blanco_pendiente = True
            continue
        if fila_en_blanco_pendiente:
            lineas.append("")
            fila_en_blanco_pendiente = False
        hubo_fila_no_vacia = True
        texto_fila = "  ".join(
            f"{get_column_letter(indice + 1)}{numero}: {texto}" for indice, texto in celdas
        )
        lineas.append(texto_fila)
    texto, advertencia = _recortar_texto_ia("\n".join(lineas), origen)
    return texto, advertencia


def _analizar_matriz_con_posible_ia(matriz, origen):
    """
    Intenta primero los caminos 100% locales (clásico y fallback por
    bloques, ambos ya existentes desde 11H). Si un header único no alcanza
    para explicar el archivo, arma un PendienteIA en vez de llamar a la IA
    desde acá — parsing.py nunca hace red ni DB; eso lo resuelve
    services.py. Devuelve (filas, advertencias, pendiente_ia | None).
    """
    if _contar_filas_header(matriz) >= 2:
        texto, advertencia_recorte = construir_texto_celdas_no_vacias(matriz, origen)
        advertencias = [advertencia_recorte] if advertencia_recorte else []
        return (
            [],
            advertencias,
            PendienteIA(
                tipo="extraccion",
                origen=origen,
                contenido=texto,
                tipo_contenido="excel_celdas",
            ),
        )

    mensaje_clasico = None
    try:
        filas, adv = _extraer_filas_matriz(matriz, origen, confianza="alta")
        return filas, adv, None
    except ColumnasNoDetectadas as exc_clasico:
        # Python borra la variable del "except ... as" al salir del bloque:
        # el mensaje se guarda acá para poder usarlo más abajo.
        mensaje_clasico = str(exc_clasico)

    filas_bloques, adv_bloques = _extraer_filas_excel_por_bloques(matriz, origen)
    candidatas = sum(1 for _, valores in matriz if _fila_parece_producto(valores))
    if filas_bloques and (candidatas == 0 or len(filas_bloques) >= candidatas * 0.5):
        return filas_bloques, adv_bloques, None

    if candidatas == 0:
        # Ni el clásico ni el fallback por bloques encontraron nada, y
        # tampoco hay señal estructural de catálogo: mismo comportamiento
        # que antes de la IA (advertencia, sin filas), sin gastar una
        # llamada que muy probablemente tampoco iba a encontrar nada.
        return filas_bloques, [mensaje_clasico, *adv_bloques], None

    posicion, mejor_fila = _mejor_candidata_encabezado(matriz)
    if mejor_fila:
        muestra_filas = [
            valores
            for _, valores in matriz[posicion + 1 : posicion + 4]
            if valores and any(_texto_celda(v) for v in valores)
        ]
        return (
            [],
            [],
            PendienteIA(
                tipo="mapeo_columnas",
                origen=origen,
                encabezados=[_texto_celda(v) for v in mejor_fila],
                muestra_filas=muestra_filas,
                matriz_restante=matriz[posicion + 1 :],
            ),
        )

    # Hay filas que parecen producto pero ni siquiera un header candidato
    # débil: último recurso antes de rendirse, extracción completa con IA.
    texto, advertencia_recorte = construir_texto_celdas_no_vacias(matriz, origen)
    advertencias = [advertencia_recorte] if advertencia_recorte else []
    return (
        [],
        advertencias,
        PendienteIA(
            tipo="extraccion",
            origen=origen,
            contenido=texto,
            tipo_contenido="excel_celdas",
        ),
    )


def _validar_zip_seguro(datos):
    try:
        with zipfile.ZipFile(BytesIO(datos)) as zf:
            total = 0
            for info in zf.infolist():
                total += info.file_size
                if total > MAX_ZIP_DESCOMPRIMIDO:
                    raise ArchivoImportacionInvalido(
                        "El archivo comprimido expande demasiado contenido y fue rechazado por seguridad."
                    )
                if info.compress_size > 0 and info.file_size / info.compress_size > MAX_RATIO_ZIP:
                    raise ArchivoImportacionInvalido(
                        "El archivo tiene una relación de compresión anormal y fue rechazado por seguridad."
                    )
    except zipfile.BadZipFile as exc:
        raise ArchivoImportacionInvalido("El archivo está dañado o no es un documento válido.") from exc


def _normalizar_imagen(contenido, origen, nombre="", numero_fila=None):
    if not contenido or len(contenido) > MAX_IMAGEN_BYTES:
        return None, "Se omitió una imagen demasiado grande."

    try:
        with warnings.catch_warnings():
            warnings.simplefilter("error", Image.DecompressionBombWarning)
            original = Image.open(BytesIO(contenido))
            ancho, alto = original.size
            if ancho * alto > MAX_IMAGEN_PIXELES:
                return None, "Se omitió una imagen con resolución excesiva."
            imagen = ImageOps.exif_transpose(original)
            imagen.thumbnail((1600, 1600))

            salida = BytesIO()
            tiene_alpha = "A" in imagen.getbands()
            if tiene_alpha:
                imagen.save(salida, format="PNG", optimize=True)
                extension = "png"
            else:
                if imagen.mode != "RGB":
                    imagen = imagen.convert("RGB")
                imagen.save(salida, format="JPEG", quality=86, optimize=True)
                extension = "jpg"
            data = salida.getvalue()
            ancho_final, alto_final = imagen.size
    except (
        Image.UnidentifiedImageError,
        Image.DecompressionBombError,
        Image.DecompressionBombWarning,
        OSError,
        ValueError,
    ):
        return None, "Se omitió una imagen embebida que no pudo validarse."

    return (
        ImagenAnalizada(
            contenido=data,
            extension=extension,
            origen=origen,
            nombre_original=nombre,
            numero_fila_origen=numero_fila,
            ancho=ancho_final,
            alto=alto_final,
            huella_sha256=hashlib.sha256(data).hexdigest(),
        ),
        None,
    )


def _agregar_imagen(resultado, contenido, origen, nombre="", numero_fila=None):
    if len(resultado.imagenes) >= MAX_IMAGENES:
        if not any("límite de imágenes" in a for a in resultado.advertencias):
            resultado.advertencias.append(
                f"Se alcanzó el límite de imágenes ({MAX_IMAGENES}); las restantes se omitieron."
            )
        return

    imagen, advertencia = _normalizar_imagen(
        contenido,
        origen=origen,
        nombre=nombre,
        numero_fila=numero_fila,
    )
    if advertencia:
        resultado.advertencias.append(f"{origen}: {advertencia}")
        return
    if imagen is None:
        return
    if any(existing.huella_sha256 == imagen.huella_sha256 for existing in resultado.imagenes):
        return
    resultado.imagenes.append(imagen)


def _analizar_xlsx(datos):
    _validar_zip_seguro(datos)
    resultado = ResultadoAnalisis()

    libro = openpyxl.load_workbook(BytesIO(datos), data_only=True, read_only=True)
    try:
        for hoja in libro.worksheets:
            matriz = []
            for numero, fila in enumerate(hoja.iter_rows(values_only=True), start=1):
                matriz.append((numero, list(fila)))
                if numero >= MAX_FILAS + 30:
                    break
            if not any(any(_texto_celda(v) for v in valores) for _, valores in matriz):
                continue
            origen = f"Hoja {hoja.title}"
            filas, adv, pendiente = _analizar_matriz_con_posible_ia(matriz, origen)
            resultado.filas.extend(filas)
            resultado.advertencias.extend(adv)
            if pendiente is not None:
                resultado.pendientes_ia.append(pendiente)
    finally:
        libro.close()

    # openpyxl no expone imágenes en read_only; se abre una segunda vez.
    try:
        libro_img = openpyxl.load_workbook(BytesIO(datos), data_only=True, read_only=False)
        try:
            for hoja in libro_img.worksheets:
                for imagen in getattr(hoja, "_images", []):
                    anchor = getattr(imagen, "anchor", None)
                    desde = getattr(anchor, "_from", None)
                    numero_fila = getattr(desde, "row", None)
                    if numero_fila is not None:
                        numero_fila += 1
                    try:
                        contenido = imagen._data()
                    except Exception:
                        resultado.advertencias.append(
                            f"Hoja {hoja.title}: no se pudo extraer una imagen embebida."
                        )
                        continue
                    _agregar_imagen(
                        resultado,
                        contenido,
                        origen=f"Hoja {hoja.title}",
                        nombre=getattr(imagen, "path", "") or "",
                        numero_fila=numero_fila,
                    )
        finally:
            libro_img.close()
    except Exception:
        resultado.advertencias.append(
            "El contenido tabular del Excel se analizó, pero no se pudieron inspeccionar sus imágenes."
        )

    return resultado


def _analizar_xls(datos):
    resultado = ResultadoAnalisis()
    try:
        libro = xlrd.open_workbook(file_contents=datos, on_demand=True)
    except xlrd.XLRDError as exc:
        raise ArchivoImportacionInvalido("El archivo .xls está dañado o no es válido.") from exc

    try:
        for hoja in libro.sheets():
            matriz = [
                (numero + 1, hoja.row_values(numero))
                for numero in range(min(hoja.nrows, MAX_FILAS + 30))
            ]
            if not matriz:
                continue
            origen = f"Hoja {hoja.name}"
            filas, adv, pendiente = _analizar_matriz_con_posible_ia(matriz, origen)
            resultado.filas.extend(filas)
            resultado.advertencias.extend(adv)
            if pendiente is not None:
                resultado.pendientes_ia.append(pendiente)
    finally:
        libro.release_resources()

    resultado.advertencias.append(
        "El formato .xls antiguo no permite extraer imágenes de forma confiable; se analizó solo la tabla."
    )
    return resultado


def _decodificar_csv(datos):
    for encoding in ("utf-8-sig", "utf-8", "cp1252", "latin-1"):
        try:
            return datos.decode(encoding), encoding
        except UnicodeDecodeError:
            continue
    raise ArchivoImportacionInvalido("No se pudo interpretar la codificación del CSV.")


def _analizar_csv(datos):
    resultado = ResultadoAnalisis()
    texto, encoding = _decodificar_csv(datos)
    muestra = texto[:8192]
    try:
        delimitador = csv.Sniffer().sniff(muestra, delimiters=",;\t|").delimiter
    except csv.Error:
        delimitador = ";"

    lector = csv.reader(texto.splitlines(), delimiter=delimitador)
    matriz = []
    for numero, fila in enumerate(lector, start=1):
        if numero > MAX_FILAS + 30:
            resultado.advertencias.append(
                f"CSV: se alcanzó el límite de {MAX_FILAS} filas."
            )
            break
        matriz.append((numero, fila))

    filas, adv, pendiente = _analizar_matriz_con_posible_ia(matriz, "CSV")
    resultado.filas.extend(filas)
    resultado.advertencias.extend(adv)
    if pendiente is not None:
        resultado.pendientes_ia.append(pendiente)
    if encoding not in ("utf-8-sig", "utf-8"):
        resultado.advertencias.append(
            f"CSV: se interpretó usando codificación {encoding}."
        )
    return resultado


def _analizar_docx(datos):
    _validar_zip_seguro(datos)
    resultado = ResultadoAnalisis()
    try:
        documento = Document(BytesIO(datos))
    except Exception as exc:
        raise ArchivoImportacionInvalido("El Word está dañado o no es un .docx válido.") from exc

    tablas_detectadas = 0
    for indice, tabla in enumerate(documento.tables, start=1):
        matriz = [
            (numero, [celda.text.strip() for celda in fila.cells])
            for numero, fila in enumerate(tabla.rows, start=1)
        ]
        if not matriz:
            continue
        origen = f"Tabla {indice} de Word"
        try:
            filas, adv = _extraer_filas_matriz(matriz, origen, confianza="alta")
        except ColumnasNoDetectadas as exc:
            resultado.advertencias.append(str(exc))
            continue
        tablas_detectadas += 1
        resultado.filas.extend(filas)
        resultado.advertencias.extend(adv)

    # Algunos Word contienen una lista tabulada en párrafos en vez de tabla.
    if tablas_detectadas == 0:
        matriz = []
        for numero, parrafo in enumerate(documento.paragraphs, start=1):
            texto = parrafo.text.strip()
            if not texto:
                continue
            celdas = [p.strip() for p in re.split(r"\t+|\s{2,}", texto) if p.strip()]
            if len(celdas) >= 2:
                matriz.append((numero, celdas))
        if matriz:
            try:
                filas, adv = _extraer_filas_matriz(
                    matriz,
                    "Párrafos tabulados de Word",
                    confianza="media",
                )
                resultado.filas.extend(filas)
                resultado.advertencias.extend(adv)
            except ColumnasNoDetectadas as exc:
                resultado.advertencias.append(str(exc))

    # Extrae imágenes desde las relaciones del paquete DOCX.
    for rel in documento.part.rels.values():
        if "image" not in rel.reltype:
            continue
        try:
            blob = rel.target_part.blob
            nombre = Path(str(rel.target_ref)).name
        except Exception:
            continue
        _agregar_imagen(
            resultado,
            blob,
            origen="Documento Word",
            nombre=nombre,
        )

    if not resultado.filas:
        texto = _texto_docx_para_ia(documento)
        if texto.strip():
            texto, advertencia_recorte = _recortar_texto_ia(texto, "Documento Word")
            if advertencia_recorte:
                resultado.advertencias.append(advertencia_recorte)
            resultado.pendientes_ia.append(
                PendienteIA(
                    tipo="extraccion",
                    origen="Documento Word",
                    contenido=texto,
                    tipo_contenido="docx_texto",
                )
            )

    return resultado


def _texto_docx_para_ia(documento):
    """
    Texto plano de tablas + párrafos para mandarle a Sonnet cuando ninguna
    tabla local fue interpretable (o no había tablas). Nunca se manda el
    .docx original: python-docx ya extrajo el texto acá, localmente.
    """
    partes = []
    for indice, tabla in enumerate(documento.tables, start=1):
        filas_tabla = []
        for fila in tabla.rows:
            celdas = [celda.text.strip() for celda in fila.cells]
            if any(celdas):
                filas_tabla.append(" | ".join(celdas))
        if filas_tabla:
            partes.append(f"Tabla {indice}:")
            partes.extend(filas_tabla)
            partes.append("")

    parrafos = [p.text.strip() for p in documento.paragraphs if p.text.strip()]
    if parrafos:
        partes.append("Texto del documento:")
        partes.extend(parrafos)

    return "\n".join(partes)


def _lineas_pdf_a_matriz(texto):
    matriz = []
    for numero, linea in enumerate((texto or "").splitlines(), start=1):
        limpia = linea.strip()
        if not limpia:
            continue
        celdas = [
            parte.strip()
            for parte in re.split(r"\t+|\s{2,}", limpia)
            if parte.strip()
        ]
        if len(celdas) >= 2:
            matriz.append((numero, celdas))
    return matriz


def _analizar_pdf(datos):
    resultado = ResultadoAnalisis()
    try:
        reader = PdfReader(BytesIO(datos), strict=False)
    except Exception as exc:
        raise ArchivoImportacionInvalido("El PDF está dañado o no se puede leer.") from exc

    if len(reader.pages) > MAX_PAGINAS_PDF:
        raise ArchivoImportacionInvalido(
            f"El PDF tiene {len(reader.pages)} páginas; el máximo permitido es {MAX_PAGINAS_PDF}."
        )

    paginas_con_texto = 0
    paginas_con_tabla = 0

    for numero_pagina, pagina in enumerate(reader.pages, start=1):
        try:
            texto = pagina.extract_text(extraction_mode="layout") or ""
        except Exception:
            try:
                texto = pagina.extract_text() or ""
            except Exception:
                texto = ""

        if texto.strip():
            paginas_con_texto += 1
            matriz = _lineas_pdf_a_matriz(texto)
            if matriz:
                origen = f"Página {numero_pagina} del PDF"
                try:
                    filas, adv = _extraer_filas_matriz(
                        matriz,
                        origen,
                        confianza="media",
                    )
                    if filas:
                        paginas_con_tabla += 1
                    resultado.filas.extend(filas)
                    resultado.advertencias.extend(adv)
                except ColumnasNoDetectadas as exc:
                    resultado.advertencias.append(str(exc))
        else:
            resultado.advertencias.append(
                f"Página {numero_pagina} del PDF: no contiene texto extraíble. "
                "Puede ser una página escaneada; se preservan sus imágenes, pero 11H "
                "no hace OCR automático para evitar interpretar precios erróneos."
            )

        try:
            imagenes_pagina = list(pagina.images)
        except Exception:
            imagenes_pagina = []
        for imagen in imagenes_pagina:
            try:
                contenido = imagen.data
                nombre = getattr(imagen, "name", "") or ""
            except Exception:
                continue
            _agregar_imagen(
                resultado,
                contenido,
                origen=f"Página {numero_pagina} del PDF",
                nombre=nombre,
            )

    if paginas_con_texto == 0:
        resultado.advertencias.append(
            "El PDF no tiene texto extraíble. Se conservaron las imágenes detectables "
            "para revisión, pero no se generaron productos automáticamente."
        )
    elif paginas_con_tabla == 0:
        resultado.advertencias.append(
            "Se pudo leer texto del PDF, pero no se detectó una tabla con Código, "
            "Nombre/Descripción y Costo/Precio."
        )

    if not resultado.filas:
        # Cubre tanto el PDF escaneado (paginas_con_texto == 0, necesita
        # visión) como el PDF con texto pero sin tabla reconocible por el
        # parser local: en los dos casos, se manda el PDF completo (nunca
        # solo texto suelto) como bloque "document" nativo, para que Sonnet
        # pueda usar layout/visión si hace falta.
        resultado.pendientes_ia.append(
            PendienteIA(
                tipo="extraccion",
                origen="PDF completo",
                contenido=datos,
                tipo_contenido="pdf",
            )
        )

    return resultado


def _analizar_imagen(datos):
    """
    Foto de una lista de precios. No hay camino local para esto (a
    diferencia de Excel/CSV/PDF/Word, nunca hubo un parser que "lea" una
    imagen como datos) — siempre se resuelve con Sonnet. La imagen
    normalizada se guarda además como evidencia visual, igual que las
    imágenes embebidas de los otros formatos.
    """
    resultado = ResultadoAnalisis()
    imagen, advertencia = _normalizar_imagen(datos, origen="Imagen cargada")
    if advertencia:
        raise ArchivoImportacionInvalido(advertencia)

    resultado.imagenes.append(imagen)
    resultado.pendientes_ia.append(
        PendienteIA(
            tipo="extraccion",
            origen="Imagen cargada",
            contenido=imagen.contenido,
            tipo_contenido="imagen",
            extension_imagen=imagen.extension,
        )
    )
    return resultado


def analizar_archivo(archivo, tipo_archivo=None):
    """
    Punto único de análisis 11H. Lee como máximo 30 MB y despacha por formato.

    Devuelve filas + imágenes + advertencias. Ningún parser toca catálogo,
    proveedores ni precios.
    """
    if tipo_archivo is None:
        tipo_archivo = tipo_archivo_por_nombre(getattr(archivo, "name", ""))

    try:
        archivo.open("rb")
    except AttributeError:
        pass
    try:
        if hasattr(archivo, "seek"):
            archivo.seek(0)
        datos = archivo.read(MAX_ARCHIVO_BYTES + 1)
    finally:
        try:
            archivo.close()
        except Exception:
            pass

    if len(datos) > MAX_ARCHIVO_BYTES:
        raise ArchivoImportacionInvalido(
            f"El archivo supera el máximo de {MAX_ARCHIVO_BYTES // (1024 * 1024)} MB."
        )

    parsers = {
        "xlsx": _analizar_xlsx,
        "xls": _analizar_xls,
        "csv": _analizar_csv,
        "pdf": _analizar_pdf,
        "docx": _analizar_docx,
        "imagen": _analizar_imagen,
    }
    parser = parsers.get(tipo_archivo)
    if parser is None:
        raise ArchivoImportacionInvalido("Formato de archivo no soportado.")

    resultado = parser(datos)
    if len(resultado.filas) > MAX_FILAS:
        resultado.filas = resultado.filas[:MAX_FILAS]
        resultado.advertencias.append(
            f"Se recortó el análisis al máximo de {MAX_FILAS} filas."
        )
    return resultado


def leer_excel(archivo):
    """
    Compatibilidad con la API de Etapa 4. Solo devuelve filas del primer
    análisis XLSX; el código nuevo usa analizar_archivo().
    """
    resultado = analizar_archivo(archivo, "xlsx")
    if not resultado.filas:
        raise ColumnasNoDetectadas(
            list(CAMPOS_REQUERIDOS),
            [],
            origen="Excel",
        )
    for fila in resultado.filas:
        yield fila.numero, {
            "marca": fila.datos["marca"],
            "codigo": fila.datos["codigo"],
            "nombre": fila.datos["nombre"],
            "costo_crudo": fila.datos["costo_crudo"],
            "codigo_proveedor": fila.datos["codigo_proveedor"],
        }
